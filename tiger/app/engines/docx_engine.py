"""
DOCX Engine for Word Document Processing
Native Word document text extraction engine
"""

import sys
import subprocess
from typing import Dict, Any
try:
    from .base_engine import BaseEngine, ExtractionResult
except ImportError:
    from base_engine import BaseEngine, ExtractionResult

class DocxEngine(BaseEngine):
    """DOCX processing engine for Word documents"""
    
    def __init__(self):
        super().__init__("DocxEngine")
        self.supported_formats = ['.docx']
        self._python_docx_available = None
    
    def setup_dependencies(self) -> bool:
        """Install and setup python-docx dependencies"""
        if self._python_docx_available is True:
            return True
        
        try:
            # Try to import python-docx
            import docx
            self._python_docx_available = True
            self.logger.info("python-docx already available")
            return True
        except ImportError:
            pass
        
        # python-docx not available - do not try to install at runtime
        self.logger.warning("python-docx not available. Install with: pip install python-docx")
        self._python_docx_available = False
        return False
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX using python-docx"""
        # Ensure dependencies are available
        if not self.setup_dependencies():
            return ExtractionResult(
                success=False,
                error="python-docx dependencies not available",
                engine_name=self.name
            )
        
        try:
            from docx import Document
            
            # Open document
            doc = Document(file_path)
            
            # Extract paragraphs
            text_content = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    paragraph_count += 1
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n".join(text_content)
            
            # Prepare metadata
            metadata = {
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'format': 'docx',
                'extraction_method': 'python_docx'
            }
            
            # Extract document properties if available
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                })
            except Exception:
                pass  # Document properties not available
            
            # Check if we got meaningful content
            if not full_text or len(full_text.strip()) < 5:
                return ExtractionResult(
                    success=False,
                    error="No meaningful text extracted from document",
                    engine_name=self.name,
                    metadata=metadata
                )
            
            return ExtractionResult(
                success=True,
                text=full_text,
                metadata=metadata,
                engine_name=self.name
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"DOCX extraction failed: {str(e)}",
                engine_name=self.name
            )
    
    def get_engine_info(self) -> Dict[str, Any]:
        """Get detailed engine information"""
        info = super().get_engine_info()
        info.update({
            'description': 'Native Word document text extraction engine',
            'features': [
                'Direct DOCX format support',
                'Paragraph structure preservation',
                'Table content extraction',
                'Document metadata extraction',
                'High-speed processing'
            ],
            'optimal_for': [
                'Attorney notes and memos',
                'Client correspondence',
                'Legal briefs and pleadings',
                'Contract documents',
                'Internal case files'
            ],
            'python_docx_available': self._python_docx_available
        })
        return info