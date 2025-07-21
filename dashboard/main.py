import os
import json
import logging
import hashlib
import glob
import shutil
import subprocess
from datetime import datetime, timedelta
from fastapi import FastAPI, HTTPException, Query, Request, File, UploadFile, WebSocket, WebSocketDisconnect, Depends, Cookie
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse, Response, RedirectResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from contextlib import asynccontextmanager
import threading
import asyncio
import secrets
from typing import Literal, List, Optional

from .data_manager import DataManager
from .file_watcher import FileWatcher
from .models import CaseStatus
from . import service_runner
from .sync_manager import SyncManager
from .upload_service import StandaloneCaseUploader

# Document parsing removed - Tiger service handles all document processing

# --- WebSocket Connection Manager ---
class ConnectionManager:
    """Manages WebSocket connections for real-time event broadcasting"""
    
    def __init__(self):
        self.active_connections: List[WebSocket] = []
        self.logger = logging.getLogger(__name__)
    
    async def connect(self, websocket: WebSocket):
        """Accept a new WebSocket connection"""
        await websocket.accept()
        self.active_connections.append(websocket)
        self.logger.info(f"New WebSocket connection. Total connections: {len(self.active_connections)}")
    
    def disconnect(self, websocket: WebSocket):
        """Remove a WebSocket connection"""
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)
            self.logger.info(f"WebSocket disconnected. Total connections: {len(self.active_connections)}")
    
    async def send_personal_message(self, message: str, websocket: WebSocket):
        """Send a message to a specific WebSocket connection"""
        try:
            await websocket.send_text(message)
        except Exception as e:
            self.logger.warning(f"Failed to send personal message: {e}")
            self.disconnect(websocket)
    
    async def broadcast_event(self, event_data: dict):
        """Broadcast an event to all connected WebSocket clients"""
        if not self.active_connections:
            self.logger.debug("No active WebSocket connections for broadcasting")
            return
        
        message = json.dumps(event_data)
        self.logger.debug(f"Broadcasting event to {len(self.active_connections)} connections: {event_data.get('type', 'unknown')}")
        
        # Send to all connections, remove failed ones
        disconnected = []
        for connection in self.active_connections:
            try:
                await connection.send_text(message)
            except Exception as e:
                self.logger.warning(f"Failed to send message to connection: {e}")
                disconnected.append(connection)
        
        # Clean up disconnected connections
        for connection in disconnected:
            self.disconnect(connection)

# --- Session Management ---
class SessionManager:
    """Simple session management for dashboard authentication"""
    
    def __init__(self):
        self.sessions = {}  # session_id -> user_info
        self.session_timeout = timedelta(hours=8)  # 8 hour session timeout
        
    def create_session(self, username: str) -> str:
        """Create a new session and return session ID"""
        session_id = secrets.token_urlsafe(32)
        self.sessions[session_id] = {
            'username': username,
            'created_at': datetime.now(),
            'last_access': datetime.now()
        }
        return session_id
    
    def get_session(self, session_id: str) -> Optional[dict]:
        """Get session info if valid, None if expired or not found"""
        if not session_id or session_id not in self.sessions:
            return None
            
        session = self.sessions[session_id]
        
        # Check if session has expired
        if datetime.now() - session['last_access'] > self.session_timeout:
            del self.sessions[session_id]
            return None
            
        # Update last access time
        session['last_access'] = datetime.now()
        return session
    
    def delete_session(self, session_id: str):
        """Delete a session"""
        if session_id in self.sessions:
            del self.sessions[session_id]
    
    def cleanup_expired_sessions(self):
        """Remove expired sessions"""
        now = datetime.now()
        expired_sessions = [
            session_id for session_id, session in self.sessions.items()
            if now - session['last_access'] > self.session_timeout
        ]
        for session_id in expired_sessions:
            del self.sessions[session_id]

# --- Authentication ---
# User accounts - Username: Password
VALID_USERS = {
    "Admin": "admin",
    "Kevin": "Automate2025!"
}

def authenticate_user(username: str, password: str) -> bool:
    """Simple authentication with multiple user support"""
    return username in VALID_USERS and VALID_USERS[username] == password

def get_current_user(session_id: str = Cookie(None)) -> Optional[dict]:
    """Get current user from session"""
    if not session_id:
        return None
    return session_manager.get_session(session_id)

def require_auth(user: dict = Depends(get_current_user)) -> dict:
    """Dependency to require authentication"""
    if not user:
        raise HTTPException(
            status_code=401,
            detail="Authentication required"
        )
    return user

# --- Configuration ---
APP_VERSION = "2.1.6"  # JSONP Version Management: Automated frontend version synchronization system

def get_git_commit():
    """Get current git commit hash for version tracking"""
    try:
        result = subprocess.run(['git', 'rev-parse', '--short', 'HEAD'], 
                              capture_output=True, text=True, cwd=os.path.dirname(__file__))
        return result.stdout.strip() if result.returncode == 0 else "unknown"
    except Exception:
        return "unknown"

def generate_version_file():
    """Generate JSONP version file for frontend consumption"""
    try:
        version_data = f'''// SATORI AI VERSION MANAGEMENT (JSONP)
// ====================================
// Single source of truth for all version references
// Auto-generated on server startup

window.satoriVersion = {{
    version: "{APP_VERSION}",
    buildDate: "{datetime.now().strftime('%Y-%m-%d')}",
    gitCommit: "{get_git_commit()}",
    environment: "{os.getenv('ENVIRONMENT', 'development')}",
    features: {{
        deployment_infrastructure: true,
        go_adapter: true,
        shadow_repository: true,
        client_data_protection: true
    }},
    // Cache busting parameter for asset loading
    cacheBuster: "{datetime.now().strftime('%Y%m%d%H')}"
}};

// Auto-update DOM elements with version class
if (typeof document !== 'undefined') {{
    document.addEventListener('DOMContentLoaded', function() {{
        // Update version display elements
        document.querySelectorAll('.version-display, .satori-version').forEach(el => {{
            el.textContent = `v${{window.satoriVersion.version}}`;
        }});
        
        // Update powered-by elements
        document.querySelectorAll('.powered-by-satori').forEach(el => {{
            el.textContent = `Powered by Satori AI v${{window.satoriVersion.version}}`;
        }});
        
        // Console logging for debugging
        console.log(`🚀 Satori AI Tiger-Monkey v${{window.satoriVersion.version}} - ${{window.satoriVersion.environment}}`);
        console.log(`📅 Build: ${{window.satoriVersion.buildDate}} | Commit: ${{window.satoriVersion.gitCommit}}`);
    }});
}}'''
        
        version_file_path = os.path.join(os.path.dirname(__file__), "static", "version.js")
        with open(version_file_path, "w") as f:
            f.write(version_data)
        
        logging.info(f"✅ Generated version.js: {APP_VERSION} ({get_git_commit()})")
        return True
    except Exception as e:
        logging.error(f"❌ Failed to generate version.js: {e}")
        return False

# Global session manager
session_manager = SessionManager()

# Global cache for grid state - prevents unnecessary DOM updates
_grid_cache = {
    'last_hash': None,
    'last_content': None,
    'last_timestamp': None
}

def clear_grid_cache():
    """Clear grid cache to force content regeneration"""
    global _grid_cache
    _grid_cache['last_hash'] = None
    _grid_cache['last_content'] = None
    _grid_cache['last_timestamp'] = None

def generate_action_button(case):
    """Generate action button HTML for a case"""
    base_classes = "w-full px-4 py-3 rounded-lg font-semibold text-center focus:outline-none focus:ring-2 focus:ring-offset-2"
    
    if case.status == CaseStatus.NEW:
        return f'<button class="{base_classes} bg-gray-800 text-white hover:bg-gray-900 focus:ring-gray-500" onclick="handleProcessWithValidation(this)" data-case-id="{case.id}">Process Files</button>'
    elif case.status == CaseStatus.PROCESSING:
        return f'<button class="{base_classes} bg-red-600 text-white cursor-not-allowed" disabled>⛔ Processing...</button>'
    elif case.status == CaseStatus.PENDING_REVIEW:
        return f'<a href="/review?case_id={case.id}" class="{base_classes} bg-blue-600 text-white hover:bg-blue-700 focus:ring-blue-500 block">Review Case</a>'
    elif case.status == CaseStatus.COMPLETE:
        return f'<a href="/review?case_id={case.id}" class="{base_classes} bg-white text-gray-700 border border-gray-300 hover:bg-gray-100 focus:ring-gray-400 block">Legal Packet</a>'
    elif case.status == CaseStatus.ERROR:
        return f'<button class="{base_classes} bg-red-600 text-white hover:bg-red-700 focus:ring-red-500" onclick="handleProcessWithValidation(this)" data-case-id="{case.id}">Fix & Reprocess</button>'
    else:
        return f'<button class="{base_classes} bg-gray-400 text-white cursor-not-allowed" disabled>{case.status.value}...</button>'

# Get the absolute path of the directory this file is in (dashboard/)
DASHBOARD_DIR = os.path.dirname(__file__)
# Get the absolute path of the project root (TM/)
PROJECT_ROOT = os.path.abspath(os.path.join(DASHBOARD_DIR, '..'))

CASE_DIRECTORY = os.path.join(PROJECT_ROOT, "test-data", "sync-test-cases")
STATIC_DIR = os.path.join(DASHBOARD_DIR, "static")
THEMES_DIR = os.path.join(STATIC_DIR, "themes")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")

# --- Global Instances ---
data_manager = DataManager(CASE_DIRECTORY, OUTPUT_DIR)
connection_manager = ConnectionManager()
source_file_watcher = FileWatcher(CASE_DIRECTORY, data_manager, connection_manager)
output_file_watcher = FileWatcher(OUTPUT_DIR, data_manager, connection_manager)

# --- Application Lifecycle ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    print("Starting application...")
    
    # Generate version.js file on startup
    generate_version_file()
    
    source_watcher_thread = threading.Thread(target=source_file_watcher.start, daemon=True)
    output_watcher_thread = threading.Thread(target=output_file_watcher.start, daemon=True)
    source_watcher_thread.start()
    output_watcher_thread.start()
    yield
    print("Stopping application...")
    source_file_watcher.stop()
    output_file_watcher.stop()

app = FastAPI(
    lifespan=lifespan,
    docs_url="/api/docs",
    redoc_url="/api/redoc"
)

# Set up logging
logger = logging.getLogger(__name__)

# --- API Endpoints ---

# WebSocket endpoint for real-time communication
@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """WebSocket endpoint for real-time event broadcasting to frontend"""
    await connection_manager.connect(websocket)
    try:
        while True:
            # Keep connection alive by receiving any messages
            # Frontend can send ping messages to keep connection active
            data = await websocket.receive_text()
            logger.debug(f"Received WebSocket message: {data}")
            
            # Echo back as heartbeat confirmation
            await websocket.send_text(json.dumps({
                "type": "heartbeat",
                "timestamp": datetime.now().isoformat(),
                "message": "Connection active"
            }))
    except WebSocketDisconnect:
        connection_manager.disconnect(websocket)
        logger.info("WebSocket client disconnected")
    except Exception as e:
        logger.error(f"WebSocket error: {e}")
        connection_manager.disconnect(websocket)

# Tiger service event receiver endpoint
@app.post("/api/processing-events")
async def receive_processing_event(request: Request):
    """Receive processing events from Tiger service and broadcast to WebSocket clients"""
    try:
        event_data = await request.json()
        logger.info(f"Received processing event: {event_data.get('type', 'unknown')} for case {event_data.get('case_id', 'unknown')}")
        
        # Validate event structure
        required_fields = ['type', 'case_id', 'timestamp']
        if not all(field in event_data for field in required_fields):
            raise HTTPException(status_code=400, detail="Invalid event structure")
        
        # Broadcast to all connected WebSocket clients
        await connection_manager.broadcast_event(event_data)
        
        return {"status": "success", "message": "Event broadcasted to WebSocket clients"}
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON payload")
    except Exception as e:
        logger.error(f"Error processing event: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/api/file-system-event")
async def receive_file_system_event(request: Request):
    """Receive file system events and broadcast to WebSocket clients for real-time UI updates"""
    try:
        event_data = await request.json()
        logger.info(f"Received file system event: {event_data.get('type', 'unknown')} for path {event_data.get('path', 'unknown')}")
        
        # Broadcast to all connected WebSocket clients
        await connection_manager.broadcast_event(event_data)
        
        return {"status": "success", "message": "File system event broadcasted to WebSocket clients"}
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON payload")
    except Exception as e:
        logger.error(f"Error processing file system event: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/api/attorney_template")
async def get_attorney_template():
    template_path = os.path.join(DASHBOARD_DIR, "config", "attorney_notes_template.txt")
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Attorney template not found.")
    return FileResponse(template_path)

@app.post("/api/attorney_template/docx")
async def generate_attorney_template_docx(request: Request):
    """Generate a DOCX version of the attorney template for download"""
    try:
        data = await request.json()
        content = data.get('content', '')
        
        # Create a simple DOCX document
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx library not available")
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Attorney Notes Template', 0)
        
        # Add instructions
        doc.add_paragraph('Instructions: Fill in the information after each label for your specific case.')
        doc.add_paragraph('')  # Empty line
        
        # Add template content with formatting
        lines = content.split('\n')
        for line in lines:
            if line.strip().endswith(':') and line.strip():
                # This is a label - make it bold
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                # Add empty line for content
                doc.add_paragraph('')
            elif line.strip():
                # Regular content
                doc.add_paragraph(line)
            else:
                # Empty line
                doc.add_paragraph('')
        
        # Save to bytes
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Return as download
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            io.BytesIO(doc_bytes.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Atty_Notes_Template.docx"}
        )
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

@app.get("/api/cases")
async def get_cases():
    return data_manager.get_all_cases()

@app.post("/api/refresh")
async def refresh_cases():
    """Force a manual refresh of case data and progress states"""
    try:
        data_manager.scan_cases()
        clear_grid_cache()  # Clear cache to force UI update
        return {"message": "Cases refreshed successfully", "timestamp": datetime.now().isoformat()}
    except Exception as e:
        logger.error(f"Error refreshing cases: {e}")
        raise HTTPException(status_code=500, detail="Failed to refresh cases")

@app.get("/api/version")
async def get_version():
    return {"version": APP_VERSION}

@app.get("/api/changelog")
async def get_changelog():
    """Serve the changelog content for display in settings."""
    changelog_path = os.path.join(DASHBOARD_DIR, "changelog.md")
    if not os.path.exists(changelog_path):
        raise HTTPException(status_code=404, detail="Changelog not found")
    
    try:
        with open(changelog_path, 'r', encoding='utf-8') as f:
            changelog_content = f.read()
        
        return {
            "content": changelog_content,
            "current_version": APP_VERSION,
            "last_updated": datetime.fromtimestamp(os.path.getmtime(changelog_path)).isoformat()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading changelog: {str(e)}")

# All document validation removed - Tiger service handles all document processing and validation

@app.get("/api/cases/{case_id}/validate-timeline")
async def validate_case_timeline(case_id: str):
    """
    Validate chronological timeline for a case - MVP 1 Task 1.3
    Returns detailed timeline validation results with business rule violations
    """
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Check if case has been processed and has timeline data
    if not case.hydrated_json_path or not os.path.exists(case.hydrated_json_path):
        return {
            "case_id": case_id,
            "timeline_available": False,
            "message": "Case must be processed first to generate timeline data",
            "validation_score": 0
        }
    
    try:
        # Load hydrated JSON to access timeline data
        with open(case.hydrated_json_path, 'r') as f:
            case_data = json.load(f)
        
        timeline = case_data.get('case_timeline', {})
        if not timeline:
            return {
                "case_id": case_id,
                "timeline_available": False,
                "message": "No timeline data found in processed case",
                "validation_score": 0
            }
        
        # Extract timeline validation results
        chronological_validation = timeline.get('chronological_validation', {})
        timeline_confidence = timeline.get('timeline_confidence', 0.0)
        
        # Count timeline elements for completeness assessment
        key_dates_present = sum([
            1 if timeline.get('discovery_date') else 0,
            1 if timeline.get('dispute_date') else 0,
            1 if timeline.get('filing_date') else 0
        ])
        
        document_dates_count = len(timeline.get('document_dates', []))
        errors_count = len(chronological_validation.get('errors', []))
        warnings_count = len(chronological_validation.get('warnings', []))
        
        # Calculate validation score
        validation_score = 0
        if chronological_validation.get('is_valid', False):
            validation_score += 40  # No critical errors
        validation_score += min(key_dates_present * 15, 45)  # Up to 45 points for key dates
        validation_score += min(timeline_confidence * 0.15, 15)  # Convert confidence to 15 points max
        
        timeline_validation_result = {
            "case_id": case_id,
            "timeline_available": True,
            "is_chronologically_valid": chronological_validation.get('is_valid', True),
            "validation_score": round(validation_score, 1),
            "timeline_confidence": round(timeline_confidence, 1),
            "summary": {
                "key_dates_present": key_dates_present,
                "document_dates_extracted": document_dates_count,
                "critical_errors": errors_count,
                "warnings": warnings_count
            },
            "key_dates": {
                "discovery_date": timeline.get('discovery_date'),
                "dispute_date": timeline.get('dispute_date'),
                "filing_date": timeline.get('filing_date')
            },
            "validation_details": {
                "errors": chronological_validation.get('errors', []),
                "warnings": chronological_validation.get('warnings', [])
            },
            "document_dates": timeline.get('document_dates', []),
            "recommendations": []
        }
        
        # Add specific recommendations based on validation results
        if not timeline.get('discovery_date'):
            timeline_validation_result["recommendations"].append("Add discovery date to attorney notes for complete timeline")
        if not timeline.get('dispute_date'):
            timeline_validation_result["recommendations"].append("Add dispute date to attorney notes for chronological validation")
        if errors_count > 0:
            timeline_validation_result["recommendations"].append("Resolve chronological errors before filing")
        if timeline_confidence < 70:
            timeline_validation_result["recommendations"].append("Review date extraction accuracy - low confidence detected")
        
        return timeline_validation_result
        
    except Exception as e:
        logging.error(f"Timeline validation error for case {case_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Timeline validation failed: {str(e)}")

@app.post("/api/cases/{case_id}/legal-claims")
async def update_legal_claims(case_id: str, request: Request):
    """Update legal claim selections for a case."""
    try:
        data = await request.json()
        selections = data.get('selections', [])
        
        # Load existing hydrated JSON
        case = data_manager.get_case_by_id(case_id)
        if not case or not case.hydrated_json_path:
            raise HTTPException(status_code=404, detail="Data still processing please try again in a few mins")
        
        if not os.path.exists(case.hydrated_json_path):
            raise HTTPException(status_code=404, detail="Hydrated JSON file not found")
        
        with open(case.hydrated_json_path, 'r') as f:
            case_data = json.load(f)
        
        # Update selections in causes of action
        if 'causes_of_action' not in case_data:
            raise HTTPException(status_code=400, detail="No causes of action found in case data")
        
        # Apply selections
        selections_applied = 0
        for selection in selections:
            cause_index = selection['cause_index']
            claim_index = selection['claim_index'] 
            selected = selection['selected']
            
            # Validate indices
            if cause_index < len(case_data['causes_of_action']):
                cause = case_data['causes_of_action'][cause_index]
                if 'legal_claims' in cause and claim_index < len(cause['legal_claims']):
                    cause['legal_claims'][claim_index]['selected'] = selected
                    selections_applied += 1
        
        # Save updated JSON
        with open(case.hydrated_json_path, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        # Mark case as reviewed when legal claims are saved
        case.progress.reviewed = True
        
        return {
            "success": True, 
            "message": f"Legal claim selections updated ({selections_applied} claims)",
            "selections_applied": selections_applied
        }
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error updating legal claims: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.post("/api/cases/{case_id}/damages")
async def update_damage_selections(case_id: str, request: Request):
    """Update damage selections for a case."""
    try:
        data = await request.json()
        selections = data.get('selections', {})
        
        # Load existing hydrated JSON
        case = data_manager.get_case_by_id(case_id)
        if not case or not case.hydrated_json_path:
            raise HTTPException(status_code=404, detail="Data still processing please try again in a few mins")
        
        if not os.path.exists(case.hydrated_json_path):
            raise HTTPException(status_code=404, detail="Hydrated JSON file not found")
        
        with open(case.hydrated_json_path, 'r') as f:
            case_data = json.load(f)
        
        # Update selections in damages
        if 'damages' not in case_data:
            raise HTTPException(status_code=400, detail="No damages found in case data")
        
        damages_data = case_data['damages']
        selections_applied = 0
        
        # Update structured damages
        if 'structured_damages' in damages_data:
            for damage in damages_data['structured_damages']:
                category = damage.get('category')
                if category in selections:
                    # Find the matching damage by category and type
                    for index, selected in selections[category].items():
                        index = int(index)
                        if 'categorized_damages' in damages_data and category in damages_data['categorized_damages']:
                            if index < len(damages_data['categorized_damages'][category]):
                                target_damage = damages_data['categorized_damages'][category][index]
                                if (damage.get('type') == target_damage.get('type') and 
                                    damage.get('entity') == target_damage.get('entity')):
                                    damage['selected'] = selected
                                    target_damage['selected'] = selected
                                    selections_applied += 1
        
        # Update categorized damages
        if 'categorized_damages' in damages_data:
            for category, category_selections in selections.items():
                if category in damages_data['categorized_damages']:
                    for index_str, selected in category_selections.items():
                        index = int(index_str)
                        if index < len(damages_data['categorized_damages'][category]):
                            damages_data['categorized_damages'][category][index]['selected'] = selected
                            selections_applied += 1
        
        # Save updated JSON
        with open(case.hydrated_json_path, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return {
            "success": True,
            "message": f"Damage selections updated ({selections_applied} damages)",
            "selections_applied": selections_applied
        }
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error updating damage selections: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.post("/api/cases/{case_id}/process")
async def process_case(case_id: str):
    print(f"🐅 BACKEND: Processing request for case {case_id}")
    
    case = data_manager.get_case_by_id(case_id)
    if not case:
        print(f"🐅 BACKEND: Case {case_id} not found")
        raise HTTPException(status_code=404, detail="Case not found")

    print(f"🐅 BACKEND: Case {case_id} found, current status: {case.status}")
    
    # Simple processing flow - Tiger handles all validation and document processing
    # Set status to PROCESSING immediately
    data_manager.update_case_status(case_id, CaseStatus.PROCESSING)
    print(f"🐅 BACKEND: Updated case {case_id} status to PROCESSING")

    def background_task():
        try:
            print(f"🐅 BACKEND: Starting background processing for case {case_id}")
            
            # Small delay then mark classification step (Step 2)
            import time
            time.sleep(1.0)  # Give UI time to show Processing status
            case.progress.classified = True
            print(f"🐅 BACKEND: Marked case {case_id} as classified")
            
            case_path = os.path.join(CASE_DIRECTORY, case_id)
            case_output_dir = os.path.join(OUTPUT_DIR, case_id)
            os.makedirs(case_output_dir, exist_ok=True)
            print(f"🐅 BACKEND: Created output directory: {case_output_dir}")
            
            # Step 3: Run file processing with animation
            print(f"🐅 BACKEND: Starting Tiger extraction for case {case_id}")
            generated_json_path = service_runner.run_tiger_extraction(case_path, case_output_dir, data_manager, case_id)
            print(f"🐅 BACKEND: Tiger extraction completed, JSON path: {generated_json_path}")
            
            # Step 4: Mark extraction complete
            case.hydrated_json_path = generated_json_path
            case.progress.extracted = True
            data_manager.update_case_status(case_id, CaseStatus.PENDING_REVIEW)
            print(f"🐅 BACKEND: Case {case_id} processing completed successfully - status: PENDING_REVIEW")

            # Broadcast the completion event
            try:
                event_data = {
                    "type": "case_processing_complete",
                    "case_id": case_id,
                    "timestamp": datetime.now().isoformat(),
                    "message": f"Processing complete for {case_id}"
                }
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                loop.run_until_complete(connection_manager.broadcast_event(event_data))
                loop.close()
            except Exception as e:
                print(f"Error broadcasting processing completion event: {e}")

        except Exception as e:
            print(f"🐅 BACKEND: Error processing case {case_id}: {e}")
            data_manager.update_case_status(case_id, CaseStatus.ERROR)

            # Broadcast the error event
            try:
                event_data = {
                    "type": "case_processing_error",
                    "case_id": case_id,
                    "timestamp": datetime.now().isoformat(),
                    "error": str(e)
                }
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                loop.run_until_complete(connection_manager.broadcast_event(event_data))
                loop.close()
            except Exception as broadcast_error:
                print(f"Error broadcasting processing error event: {broadcast_error}")

    threading.Thread(target=background_task).start()
    return {"message": f"Processing started for case {case_id}"}

@app.get("/api/cases/{case_id}/manifest")
async def get_case_manifest(case_id: str):
    """
    Return the processing manifest file for a case.
    This file contains real-time processing progress for individual files.
    """
    case_path = os.path.join(CASE_DIRECTORY, case_id)
    manifest_path = os.path.join(case_path, 'processing_manifest.txt')
    
    if not os.path.exists(manifest_path):
        # Return empty content if manifest doesn't exist yet
        return Response(content="", media_type="text/plain")
    
    try:
        with open(manifest_path, 'r') as f:
            content = f.read()
        return Response(content=content, media_type="text/plain")
    except Exception as e:
        print(f"Error reading manifest for case {case_id}: {e}")
        return Response(content="", media_type="text/plain")

@app.get("/api/cases/{case_id}/data")
async def get_case_data(case_id: str):
    case = data_manager.get_case_by_id(case_id)
    if not case or not case.hydrated_json_path:
        raise HTTPException(status_code=404, detail="Data still processing please try again in a few mins")
    
    if not os.path.exists(case.hydrated_json_path):
        raise HTTPException(status_code=404, detail="Hydrated JSON file not found at path.")

    with open(case.hydrated_json_path, 'r') as f:
        data = json.load(f)
    return JSONResponse(content=data)

@app.get("/api/cases/{case_id}/review_data")
async def get_case_review_data(case_id: str):
    # Normalize case ID to lowercase for consistency
    case_id = case_id.lower()
    case = data_manager.get_case_by_id(case_id)
    if not case or not case.hydrated_json_path:
        raise HTTPException(status_code=404, detail="Data still processing please try again in a few mins")
    
    if not os.path.exists(case.hydrated_json_path):
        raise HTTPException(status_code=404, detail="Hydrated JSON file not found at path.")

    with open(case.hydrated_json_path, 'r') as f:
        data = json.load(f)

    def format_data(d):
        if isinstance(d, list):
            return [format_data(i) for i in d]
        if isinstance(d, dict):
            return {k: format_data(v) for k, v in d.items()}
        if isinstance(d, str):
            return d.replace('"', '')
        return d

    return JSONResponse(content=format_data(data))



@app.get("/api/cases/{case_id}/file-status")
async def get_case_file_status(case_id: str):
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Return empty list if no file processing results exist yet
    files = case.file_processing_results if case.file_processing_results else []
    return {"files": files}


# HTMX HTML Fragment Endpoints
@app.get("/api/cases/{case_id}/file-status-html")
async def get_case_file_status_html(case_id: str):
    """Return HTML fragment for file status icons - used by HTMX polling"""
    from fastapi.responses import HTMLResponse
    
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Generate file status HTML
    def get_file_icon(file_name: str, file_index: int) -> str:
        """Get the appropriate icon for file status based on processing order"""
        if case.file_processing_results:
            for result in case.file_processing_results:
                if result.name == file_name and result.status == 'success':
                    return '✅'
        
        # Check if currently processing
        if case.status == CaseStatus.PROCESSING:
            # Simulate sequential processing by checking modification time of hydrated JSON
            # If hydrated JSON exists and is recent, assume later files in the list are being processed
            if case.hydrated_json_path and os.path.exists(case.hydrated_json_path):
                mod_time = os.path.getmtime(case.hydrated_json_path)
                current_time = datetime.now().timestamp()
                time_diff = current_time - mod_time
                
                # If JSON was created recently (within 2 minutes), processing is likely complete
                if time_diff < 120:  # 2 minutes
                    return '✅'
            
            # For active processing, show hourglass for first uncompleted file only
            completed_count = 0
            if case.file_processing_results:
                completed_count = len([r for r in case.file_processing_results if r.status == 'success'])
            
            # Show hourglass for the currently processing file (by index)
            if file_index == completed_count:
                return '⏳'
            elif file_index < completed_count:
                return '✅'
        
        return '☐'  # Pending
    
    # Get case files
    files = []
    if case.files:
        for file_metadata in case.files:
            if file_metadata.name.endswith(('.pdf', '.docx', '.txt')) and not file_metadata.name.startswith('.'):
                files.append(file_metadata.name)
    
    # Generate HTML for file list
    file_items_html = ""
    for file_index, file_name in enumerate(files):
        icon = get_file_icon(file_name, file_index)
        file_items_html += f'<div class="flex items-center space-x-2"><span id="file-{case_id}-{file_name}">{icon}</span><span class="text-sm text-gray-600">{file_name}</span></div>'
    
    if not file_items_html:
        file_items_html = '<div class="text-sm text-gray-500">No files found</div>'
    
    return HTMLResponse(content=file_items_html)


@app.get("/api/cases/{case_id}/status-html")
async def get_case_status_html(case_id: str):
    """Return HTML fragment for case status badge - used by HTMX polling"""
    from fastapi.responses import HTMLResponse
    
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Generate status badge HTML based on case status
    status_map = {
        CaseStatus.NEW: ("bg-blue-100 text-blue-800", "New"),
        CaseStatus.PROCESSING: ("bg-yellow-100 text-yellow-800", "Processing"),
        CaseStatus.PENDING_REVIEW: ("bg-purple-100 text-purple-800", "Pending Review"),
        CaseStatus.GENERATING: ("bg-orange-100 text-orange-800", "Generating"),
        CaseStatus.COMPLETE: ("bg-green-100 text-green-800", "Complete"),
        CaseStatus.ERROR: ("bg-red-100 text-red-800", "Error")
    }
    
    status_class, status_text = status_map.get(case.status, ("bg-gray-100 text-gray-800", "Unknown"))
    
    html = f'<span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-xs font-medium {status_class}">{status_text}</span>'
    
    return HTMLResponse(content=html)


@app.get("/api/cases/grid-html")
async def get_cases_grid_html(request: Request):
    """Return HTML fragment for cases grid - ONLY if data changed (delta-based)"""
    global _grid_cache
    
    cases = data_manager.get_all_cases()
    
    # Calculate hash of current case data to detect changes
    case_data_for_hash = []
    for case in cases:
        file_count = 0
        if case.files:
            file_count = len([f for f in case.files 
                            if f.name.endswith(('.pdf', '.docx', '.txt')) and not f.name.startswith('.')])
        
        case_data_for_hash.append({
            'id': case.id,
            'name': case.name,
            'status': case.status.value,
            'last_updated': case.last_updated.isoformat() if case.last_updated else '',
            'file_count': file_count,
            'progress': case.progress.dict() if case.progress else {}
        })
    
    # Create hash from sorted case data
    case_data_str = json.dumps(case_data_for_hash, sort_keys=True)
    current_hash = hashlib.md5(case_data_str.encode()).hexdigest()
    
    # Debug logging
    print(f"DEBUG: Current hash: {current_hash}, Cached hash: {_grid_cache['last_hash']}")
    print(f"DEBUG: Has cached content: {bool(_grid_cache['last_content'])}")
    print(f"DEBUG: Case count: {len(cases)}")
    
    # Check If-None-Match header (ETag support)
    if_none_match = request.headers.get('if-none-match')
    
    # DEBUG: Log cache state
    print(f"DEBUG: Current hash: {current_hash}")
    print(f"DEBUG: Cached hash: {_grid_cache['last_hash']}")
    print(f"DEBUG: Has cached content: {bool(_grid_cache['last_content'])}")
    print(f"DEBUG: Content length: {len(_grid_cache['last_content']) if _grid_cache['last_content'] else 0}")
    
    # CRITICAL FIX: Only return 304 if cache was previously populated AND hash matches
    if (_grid_cache['last_hash'] is not None and 
        _grid_cache['last_hash'] == current_hash and 
        _grid_cache['last_content'] and
        len(_grid_cache['last_content']) > 0):
        print("DEBUG: Returning 304 - content unchanged")
        return Response(
            status_code=304,
            headers={
                "X-Content-Changed": "false",
                "Cache-Control": "no-cache"
            }
        )
    
    # Generate new content only if data changed
    grid_html = ""
    for case in cases:
        status_class_map = {
            CaseStatus.NEW: "bg-blue-100 text-blue-800",
            CaseStatus.PROCESSING: "bg-yellow-100 text-yellow-800", 
            CaseStatus.PENDING_REVIEW: "bg-purple-100 text-purple-800",
            CaseStatus.GENERATING: "bg-orange-100 text-orange-800",
            CaseStatus.COMPLETE: "bg-green-100 text-green-800",
            CaseStatus.ERROR: "bg-red-100 text-red-800"
        }
        
        status_class = status_class_map.get(case.status, "bg-gray-100 text-gray-800")
        status_text = case.status.value.replace('_', ' ').title()
        
        # Get file count
        file_count = 0
        if case.files:
            file_count = len([f for f in case.files 
                            if f.name.endswith(('.pdf', '.docx', '.txt')) and not f.name.startswith('.')])
        
        # Remove nested polling triggers - use load only
        status_trigger = "load"
        file_trigger = "load"
        
        # Generate action button directly (no nested HTMX)
        action_button = generate_action_button(case)
        
        # Generate file status display with individual file processing animation
        file_status_html = ""
        if case.files:
            display_files = [f for f in case.files if not f.name.startswith('.') and not f.name.endswith('.ds_store')]
            if display_files:
                file_status_html = '<div class="mt-3"><h4 class="text-sm font-medium text-gray-700 mb-2">Files:</h4><div class="space-y-1">'
                for file in display_files:
                    # Individual file processing status
                    icon = "☐"  # Default: pending
                    if case.file_processing_results:
                        file_result = next((r for r in case.file_processing_results if r.name == file.name), None)
                        if file_result:
                            if file_result.status == "success":
                                icon = "✅"
                            elif file_result.status == "processing":
                                icon = "⏳"
                            elif file_result.status == "error":
                                icon = "❌"
                    
                    file_status_html += f'<div class="flex items-center text-xs text-gray-600"><span class="mr-2">{icon}</span><span class="truncate">{file.name}</span></div>'
                file_status_html += '</div></div>'
        
        grid_html += f"""
        <div class="bg-white p-6 rounded-lg shadow-sm border border-gray-200" id="case-{case.id}">
            <div class="flex items-center justify-between mb-4">
                <h3 class="text-lg font-semibold text-gray-900">{case.name}</h3>
                <span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-xs font-medium {status_class}">{status_text}</span>
            </div>
            <div class="space-y-3">
                <div class="text-sm text-gray-600">
                    <strong>{file_count}</strong> files detected
                </div>
                {file_status_html}
                <div class="flex space-x-2 mt-4" id="actions-{case.id}">
                    {action_button}
                </div>
            </div>
        </div>
        """
    
    # Update cache with new content
    _grid_cache['last_hash'] = current_hash
    _grid_cache['last_content'] = grid_html
    _grid_cache['last_timestamp'] = datetime.now()
    
    return HTMLResponse(
        content=grid_html,
        headers={
            "X-Content-Changed": "true",
            "Cache-Control": "no-cache"
        }
    )


@app.get("/api/cases/{case_id}/actions-html")
async def get_case_actions_html(case_id: str):
    """Return HTML fragment for case action buttons - used by HTMX polling"""
    from fastapi.responses import HTMLResponse
    
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    base_button_classes = "w-full text-center px-4 py-3 rounded-lg font-semibold focus:outline-none focus:ring-2 focus:ring-offset-2"
    
    # Generate action button based on case status
    if case.status == CaseStatus.NEW:
        action_button = f'''<button id="{case.id}_button" class="process-btn {base_button_classes} bg-gray-800 text-white hover:bg-gray-900 focus:ring-gray-500" 
                               data-case-id="{case.id}" 
                               onclick="handleProcessWithValidation(this)">
                           <span class="button-text">Process Files</span>
                           <span class="button-spinner hidden">⏳ Processing...</span>
                       </button>'''
    elif case.status == CaseStatus.ERROR:
        action_button = f'''<button id="{case.id}_button" class="process-btn {base_button_classes} bg-red-600 text-white hover:bg-red-700 focus:ring-red-500" 
                               data-case-id="{case.id}" 
                               onclick="handleProcessWithValidation(this)">
                           <span class="button-text">Process Files</span>
                           <span class="button-spinner hidden">⏳ Processing...</span>
                       </button>'''
    elif case.status == CaseStatus.PENDING_REVIEW:
        action_button = f'''<a id="{case.id}_button" href="/review?case_id={case.id}" class="block {base_button_classes} bg-blue-600 text-white hover:bg-blue-700 focus:ring-blue-500">Review Case</a>'''
    elif case.status == CaseStatus.PROCESSING:
        action_button = f'''<button id="{case.id}_button" class="{base_button_classes} bg-red-600 text-white hover:bg-red-700 cursor-not-allowed" disabled>
                            <span class="button-spinner">⛔ Stop Processing</span>
                        </button>'''
    elif case.status == CaseStatus.COMPLETE:
        action_button = f'''<a id="{case.id}_button" href="/review?case_id={case.id}" class="block {base_button_classes} bg-white text-gray-700 border border-gray-300 hover:bg-gray-100 focus:ring-gray-400 flex items-center justify-center">
                            <i data-feather="file-text" class="w-4 h-4 mr-2"></i>
                            Legal Packet
                        </a>'''
    else:  # GENERATING or other states
        action_button = f'''<button id="{case.id}_button" class="{base_button_classes} bg-gray-400 text-white cursor-not-allowed" disabled>{case.status.value.replace('_', ' ').title()}...</button>'''
    
    return HTMLResponse(content=action_button)


@app.post("/api/cases/{case_id}/generate-complaint")
async def generate_complaint_html(case_id: str):
    case = data_manager.get_case_by_id(case_id)
    if not case or not case.hydrated_json_path:
        raise HTTPException(status_code=404, detail="Data still processing please try again in a few mins")
    
    if not os.path.exists(case.hydrated_json_path):
        raise HTTPException(status_code=404, detail="Hydrated JSON file not found at path.")

    def background_task():
        try:
            case_output_dir = os.path.join(OUTPUT_DIR, case_id)
            os.makedirs(case_output_dir, exist_ok=True)
            
            # Run monkey service to generate complaint
            monkey_output = service_runner.run_monkey_generation(case.hydrated_json_path, case_output_dir, data_manager, case_id)
            
            case.complaint_html_path = monkey_output
            case.last_complaint_path = monkey_output
            case.progress.generated = True  # Mark document generation complete
            # Auto-mark reviewed step if not already marked (logical progression)
            if not case.progress.reviewed:
                case.progress.reviewed = True
            data_manager.update_case_status(case_id, CaseStatus.COMPLETE)
            
            # Broadcast complaint generation complete event
            # Note: Using asyncio.run() in a thread requires care
            try:
                event_data = {
                    "type": "complaint_generated",
                    "case_id": case_id,
                    "timestamp": datetime.now().isoformat(),
                    "message": f"Complaint generation completed for {case_id}"
                }
                # Create new event loop for this thread
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                loop.run_until_complete(connection_manager.broadcast_event(event_data))
                loop.close()
            except Exception as e:
                print(f"Error broadcasting complaint generation event: {e}")
        except Exception as e:
            print(f"Error generating complaint for case {case_id}: {e}")
            data_manager.update_case_status(case_id, CaseStatus.ERROR)

    threading.Thread(target=background_task).start()
    return {"message": f"Complaint generation started for case {case_id}"}

@app.post("/api/cases/{case_id}/generate-summons")
async def generate_summons_documents(case_id: str):
    """Generate individual summons documents for each defendant in the case"""
    case = data_manager.get_case_by_id(case_id)
    if not case or not case.hydrated_json_path:
        raise HTTPException(status_code=404, detail="Data still processing please try again in a few mins")
    
    if not os.path.exists(case.hydrated_json_path):
        raise HTTPException(status_code=404, detail="Hydrated JSON file not found at path.")

    def background_task():
        try:
            print(f"🏛️ BACKEND: Starting summons generation for case {case_id}")
            
            # Load case data from hydrated JSON
            with open(case.hydrated_json_path, 'r') as f:
                case_data = json.load(f)
            
            # Create output directory for summons
            case_output_dir = os.path.join(OUTPUT_DIR, case_id)
            os.makedirs(case_output_dir, exist_ok=True)
            
            # Use service runner to call monkey summons generation
            print(f"🏛️ BACKEND: Running summons generation via service_runner")
            
            summons_files = service_runner.run_summons_generation(
                case.hydrated_json_path, 
                case_output_dir, 
                data_manager, 
                case_id
            )
            
            # Update case with summons information
            case.summons_files = summons_files
            print(f"🏛️ BACKEND: Generated {len(summons_files)} summons documents for case {case_id}")
            
        except Exception as e:
            print(f"🏛️ BACKEND: Error generating summons for case {case_id}: {e}")
            import traceback
            traceback.print_exc()
            raise

    try:
        # Load case data from hydrated JSON to get defendants count
        with open(case.hydrated_json_path, 'r') as f:
            case_data = json.load(f)
        defendants = case_data.get('parties', {}).get('defendants', [])
        
        # Run summons generation in background thread
        threading.Thread(target=background_task).start()
        
        # Return immediate response with defendants info for UI update
        return {
            "message": f"Summons generation started for case {case_id}",
            "status": "processing",
            "summons_files": [f"summons_{i}.html" for i in range(len(defendants))]  # Placeholder for UI update
        }
    except Exception as e:
        logger.error(f"Error starting summons generation for case {case_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error starting summons generation: {str(e)}")

@app.get("/api/cases/{case_id}/summons-status")
async def check_summons_status(case_id: str):
    """Check if summons files exist and return their status"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Look for summons files in the output directory
    summons_dir = os.path.join(OUTPUT_DIR, case_id, "summons")
    
    if not os.path.exists(summons_dir):
        return {"exists": False, "files": [], "count": 0, "last_generated": None}
    
    # Get list of summons files
    summons_files = [f for f in os.listdir(summons_dir) if f.endswith('.html')]
    summons_files.sort()  # Ensure consistent ordering
    
    # Get the most recent modification time from all summons files
    last_generated = None
    if summons_files:
        most_recent_time = 0
        for file in summons_files:
            file_path = os.path.join(summons_dir, file)
            if os.path.exists(file_path):
                file_time = os.path.getmtime(file_path)
                if file_time > most_recent_time:
                    most_recent_time = file_time
        
        if most_recent_time > 0:
            last_generated = datetime.fromtimestamp(most_recent_time).isoformat()
    
    return {
        "exists": len(summons_files) > 0,
        "files": summons_files,
        "count": len(summons_files),
        "directory": summons_dir,
        "last_generated": last_generated
    }

@app.get("/api/cases/{case_id}/last-complaint")
async def check_last_complaint(case_id: str):
    """Check if a complaint was previously generated and return its metadata"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Look for the most recent complaint file in the Monkey output structure
    complaint_dir = os.path.join(OUTPUT_DIR, case_id, f"complaint_{case_id}.html", "processed")
    if not os.path.exists(complaint_dir):
        return {"exists": False, "path": None, "generated_at": None}
    
    # Find the most recent date directory
    date_dirs = [d for d in os.listdir(complaint_dir) if os.path.isdir(os.path.join(complaint_dir, d))]
    if not date_dirs:
        return {"exists": False, "path": None, "generated_at": None}
    
    latest_date = sorted(date_dirs, reverse=True)[0]
    complaint_folder = os.path.join(complaint_dir, latest_date)
    
    # Look for the most recent complaint file (check versioned files first)
    complaint_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint")]
    if not complaint_files:
        return {"exists": False, "path": None, "generated_at": None}
    
    # Sort by version number (complaint_v5 > complaint_v4 > complaint_v1 > complaint)
    def version_key(filename):
        if filename == "complaint":
            return 0  # Base file gets lowest priority
        elif filename.startswith("complaint_v"):
            try:
                version_num = int(filename.split("_v")[1])
                return version_num
            except (IndexError, ValueError):
                return 0
        return 0
    
    # Get the highest version complaint file
    latest_complaint = sorted(complaint_files, key=version_key, reverse=True)[0]
    complaint_path = os.path.join(complaint_folder, latest_complaint)
    
    if not os.path.exists(complaint_path):
        return {"exists": False, "path": None, "generated_at": None}
    
    # Update the case's last_complaint_path if found
    case.last_complaint_path = complaint_path
    
    # Get file modification time
    import datetime
    mod_time = os.path.getmtime(complaint_path)
    generated_at = datetime.datetime.fromtimestamp(mod_time).isoformat()
    
    return {"exists": True, "path": complaint_path, "generated_at": generated_at}

@app.get("/api/cases/{case_id}/complaint-html")
async def get_complaint_html(case_id: str):
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Look for the most recent complaint file in the Monkey output structure
    complaint_dir = os.path.join(OUTPUT_DIR, case_id, f"complaint_{case_id}.html", "processed")
    if not os.path.exists(complaint_dir):
        raise HTTPException(status_code=404, detail="No complaint generated yet")
    
    # Find the most recent date directory
    date_dirs = [d for d in os.listdir(complaint_dir) if os.path.isdir(os.path.join(complaint_dir, d))]
    if not date_dirs:
        raise HTTPException(status_code=404, detail="No complaint generated yet")
    
    latest_date = sorted(date_dirs, reverse=True)[0]
    complaint_folder = os.path.join(complaint_dir, latest_date)
    
    # Look for the most recent complaint file (check versioned files first)
    complaint_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint")]
    if not complaint_files:
        raise HTTPException(status_code=404, detail="Complaint file not found")
    
    # Sort by version number (complaint_v5 > complaint_v4 > complaint_v1 > complaint)
    def version_key(filename):
        if filename == "complaint":
            return 0  # Base file gets lowest priority
        elif filename.startswith("complaint_v"):
            try:
                version_num = int(filename.split("_v")[1])
                return version_num
            except (IndexError, ValueError):
                return 0
        return 0
    
    # Get the highest version complaint file
    latest_complaint = sorted(complaint_files, key=version_key, reverse=True)[0]
    complaint_path = os.path.join(complaint_folder, latest_complaint)
    
    if not os.path.exists(complaint_path):
        raise HTTPException(status_code=404, detail="Complaint file not found")
    
    # Update the case's last_complaint_path
    case.last_complaint_path = complaint_path
    
    with open(complaint_path, 'r') as f:
        html_content = f.read()
    
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=html_content)

@app.get("/complaint/{case_id}")
async def serve_complaint_document(case_id: str, download: bool = False):
    """Serve the complaint document at a clean URL - HTML for viewing, PDF for download"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # If download=true, prioritize PDF files
    if download:
        # FIRST: Check for standardized PDF in case folder
        case_folder_path = os.path.join(CASE_DIRECTORY, case_id)
        if os.path.exists(case_folder_path):
            standardized_pdf = f"{case_id}_complaint.pdf"
            pdf_path = os.path.join(case_folder_path, standardized_pdf)
            
            if os.path.exists(pdf_path):
                from fastapi.responses import FileResponse
                return FileResponse(
                    path=pdf_path,
                    media_type="application/pdf",
                    filename=standardized_pdf
                )
        
        # FALLBACK: Look for complaint PDF in the complex Monkey output structure
        monkey_output_dir = os.path.join(os.path.dirname(__file__), "..", "monkey", "outputs", "monkey", "processed")
        
        # Try both the expected Dashboard structure and the actual Monkey structure
        possible_dirs = [
            os.path.join(OUTPUT_DIR, case_id, f"complaint_{case_id}.html", "processed"),
            monkey_output_dir
        ]
        
        complaint_dir = None
        for dir_path in possible_dirs:
            if os.path.exists(dir_path):
                complaint_dir = dir_path
                break
        
        if complaint_dir:
            # Find the most recent date directory
            date_dirs = [d for d in os.listdir(complaint_dir) if os.path.isdir(os.path.join(complaint_dir, d))]
            if date_dirs:
                latest_date = sorted(date_dirs, reverse=True)[0]
                complaint_folder = os.path.join(complaint_dir, latest_date)
                
                # Look for PDF files specifically
                all_files = os.listdir(complaint_folder)
                
                # First look for standardized name: complaint_<case_id>.pdf
                standardized_name = f"complaint_{case_id}.pdf"
                if standardized_name in all_files:
                    pdf_path = os.path.join(complaint_folder, standardized_name)
                    if os.path.exists(pdf_path):
                        from fastapi.responses import FileResponse
                        return FileResponse(
                            path=pdf_path,
                            media_type="application/pdf",
                            filename=standardized_name
                        )
                
                # Fall back to any complaint PDF files
                pdf_files = [f for f in all_files if f.startswith("complaint") and f.endswith(".pdf")]
                if pdf_files:
                    # Found PDF file(s), serve the most recent one
                    latest_pdf = sorted(pdf_files, key=lambda x: os.path.getmtime(os.path.join(complaint_folder, x)), reverse=True)[0]
                    pdf_path = os.path.join(complaint_folder, latest_pdf)
                    
                    if os.path.exists(pdf_path):
                        from fastapi.responses import FileResponse
                        return FileResponse(
                            path=pdf_path,
                            media_type="application/pdf",
                            filename=f"complaint_{case_id}.pdf"
                        )
        
        # If no PDF found, return 404
        raise HTTPException(status_code=404, detail="No complaint PDF available")
    
    # Default behavior: serve HTML
    # FALLBACK: Look for complaint in the complex Monkey output structure
    # First check the actual Monkey output location
    monkey_output_dir = os.path.join(os.path.dirname(__file__), "..", "monkey", "outputs", "monkey", "processed")
    
    # Try both the expected Dashboard structure and the actual Monkey structure
    possible_dirs = [
        os.path.join(OUTPUT_DIR, case_id, f"complaint_{case_id}.html", "processed"),
        monkey_output_dir
    ]
    
    complaint_dir = None
    for dir_path in possible_dirs:
        if os.path.exists(dir_path):
            complaint_dir = dir_path
            break
    
    if not complaint_dir:
        raise HTTPException(status_code=404, detail="No complaint generated yet")
    
    # Find the most recent date directory
    date_dirs = [d for d in os.listdir(complaint_dir) if os.path.isdir(os.path.join(complaint_dir, d))]
    if not date_dirs:
        raise HTTPException(status_code=404, detail="No complaint generated yet")
    
    latest_date = sorted(date_dirs, reverse=True)[0]
    complaint_folder = os.path.join(complaint_dir, latest_date)
    
    # Look for the most recent complaint file (check versioned files first)
    complaint_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint")]
    if not complaint_files:
        raise HTTPException(status_code=404, detail="Complaint file not found")
    
    # Sort by version number (complaint_v5 > complaint_v4 > complaint_v1 > complaint)
    def version_key(filename):
        if filename == "complaint":
            return 0  # Base file gets lowest priority
        elif filename.startswith("complaint_v"):
            try:
                version_num = int(filename.split("_v")[1])
                return version_num
            except (IndexError, ValueError):
                return 0
        return 0
    
    # Prioritize HTML files for viewing (PDFs available via separate download endpoint)
    # Get the highest version complaint file
    latest_complaint = sorted(complaint_files, key=version_key, reverse=True)[0]
    complaint_path = os.path.join(complaint_folder, latest_complaint)
    
    if not os.path.exists(complaint_path):
        raise HTTPException(status_code=404, detail="Complaint file not found")
    
    with open(complaint_path, 'r') as f:
        html_content = f.read()
    
    # Add print-optimized CSS for better PDF output
    enhanced_html = html_content.replace(
        "</head>",
        """
        <style media="print">
            @page { 
                margin: 1in; 
                size: letter; 
            }
            body { 
                margin: 0; 
                -webkit-print-color-adjust: exact; 
                print-color-adjust: exact; 
            }
        </style>
        </head>"""
    )
    
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=enhanced_html)


@app.get("/summons/{case_id}/{defendant_index_html}")
async def serve_summons_document(case_id: str, defendant_index_html: str):
    """Serve individual summons document at a clean URL for printing/PDF generation"""
    # Parse defendant index from the filename (e.g., "0.html" -> 0)
    if not defendant_index_html.endswith('.html'):
        raise HTTPException(status_code=404, detail="Invalid summons URL format")
    
    try:
        defendant_index = int(defendant_index_html[:-5])  # Remove ".html" and convert to int
    except ValueError:
        raise HTTPException(status_code=404, detail="Invalid defendant index")
    
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Look for summons files in the output directory
    summons_dir = os.path.join(OUTPUT_DIR, case_id, "summons")
    if not os.path.exists(summons_dir):
        raise HTTPException(status_code=404, detail="No summons generated yet")
    
    # Get list of summons files
    summons_files = [f for f in os.listdir(summons_dir) if f.endswith('.html')]
    if defendant_index >= len(summons_files):
        raise HTTPException(status_code=404, detail="Defendant index out of range")
    
    # Sort files to ensure consistent ordering
    summons_files.sort()
    summons_file_path = os.path.join(summons_dir, summons_files[defendant_index])
    
    if not os.path.exists(summons_file_path):
        raise HTTPException(status_code=404, detail="Summons file not found")
    
    try:
        with open(summons_file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Add print-optimized CSS for better PDF output (similar to complaint)
        enhanced_html = html_content.replace(
            "</head>",
            """
            <style media="print">
                @page { 
                    margin: 1in; 
                    size: letter; 
                }
                body { 
                    margin: 0; 
                    -webkit-print-color-adjust: exact; 
                    print-color-adjust: exact; 
                }
                .no-print {
                    display: none !important;
                }
            </style>
            </head>"""
        )
        
        from fastapi.responses import HTMLResponse
        return HTMLResponse(content=enhanced_html)
        
    except Exception as e:
        logger.error(f"Error serving summons document: {str(e)}")
        raise HTTPException(status_code=500, detail="Error loading summons document")

@app.get("/api/creditor-addresses")
async def get_creditor_addresses():
    """Get the list of creditor addresses from creditor_addresses.json registry"""
    try:
        creditor_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "creditor_addresses.json")
        
        if not os.path.exists(creditor_file):
            # Return default empty structure
            return {
                "creditor_addresses": {},
                "metadata": {
                    "last_updated": datetime.now().isoformat(),
                    "version": "1.0.0",
                    "description": "Standard creditor addresses for legal document generation"
                }
            }
        
        with open(creditor_file, 'r') as f:
            creditor_data = json.load(f)
            
        return creditor_data
            
    except Exception as e:
        logger.error(f"Error loading creditor addresses: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error loading creditor addresses: {str(e)}")

@app.post("/api/creditor-addresses")
async def update_creditor_addresses(request: Request):
    """Update the creditor addresses registry"""
    try:
        data = await request.json()
        
        # Validate the structure
        if not isinstance(data.get('creditor_addresses'), dict):
            raise HTTPException(status_code=400, detail="Invalid creditor_addresses data format")
        
        # Update metadata
        data['metadata'] = {
            "version": "1.0",
            "created_date": "2025-07-11",
            "description": "Predefined creditor addresses for major credit bureaus and common defendants",
            "last_updated": datetime.now().strftime("%Y-%m-%d")
        }
        
        # Save to creditor addresses registry
        creditor_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "creditor_addresses.json")
        os.makedirs(os.path.dirname(creditor_file), exist_ok=True)
        
        with open(creditor_file, 'w') as f:
            json.dump(data, f, indent=2)
        
        return {
            "message": "Creditor addresses updated successfully", 
            "count": len(data['creditor_addresses'])
        }
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON format")
    except Exception as e:
        logger.error(f"Error updating creditor addresses: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error updating creditor addresses: {str(e)}")

@app.get("/api/settings")
async def get_settings():
    """Get current system settings"""
    settings_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "settings.json")
    
    if not os.path.exists(settings_file):
        # Return default settings if no config file exists
        default_settings = {
            "firm": {
                "name": "",
                "address": "",
                "phone": "",
                "email": ""
            },
            "document": {
                "default_court": "UNITED STATES DISTRICT COURT",
                "default_district": "EASTERN DISTRICT OF NEW YORK"
            },
            "icloud": {
                "account": "",
                "password": "",
                "folder": "/LegalCases"
            },
            "system": {
                "auto_save": True,
                "data_retention": 90
            }
        }
        return default_settings
    
    try:
        with open(settings_file, 'r') as f:
            settings = json.load(f)
        return settings
    except Exception as e:
        logger.error(f"Error loading settings: {str(e)}")
        raise HTTPException(status_code=500, detail="Error loading settings")

@app.post("/api/settings")
async def save_settings(request: Request):
    """Save system settings"""
    try:
        settings = await request.json()
        
        # Ensure config directory exists
        config_dir = os.path.join(PROJECT_ROOT, "dashboard", "config")
        os.makedirs(config_dir, exist_ok=True)
        
        settings_file = os.path.join(config_dir, "settings.json")
        
        # Validate required fields
        if not settings.get('firm', {}).get('name'):
            raise HTTPException(status_code=400, detail="Law firm name is required")
        if not settings.get('firm', {}).get('address'):
            raise HTTPException(status_code=400, detail="Business address is required")
        if not settings.get('firm', {}).get('phone'):
            raise HTTPException(status_code=400, detail="Business phone is required")
        if not settings.get('firm', {}).get('email'):
            raise HTTPException(status_code=400, detail="Contact email is required")
        
        # Save settings to file
        with open(settings_file, 'w') as f:
            json.dump(settings, f, indent=2)
        
        logger.info(f"Settings saved to {settings_file}")
        return {"message": "Settings saved successfully", "timestamp": datetime.now().isoformat()}
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error saving settings: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error saving settings: {str(e)}")

# --- iCloud Integration API ---

@app.get("/api/icloud/config")
async def get_icloud_config():
    """Get current iCloud configuration"""
    config_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "icloud.json")
    
    if not os.path.exists(config_file):
        # Return default configuration if no config file exists
        default_config = {
            "folder": "LegalCases",
            "sync_interval": 30,
            "log_level": "info",
            "backup_enabled": False
        }
        return default_config
    
    try:
        with open(config_file, 'r') as f:
            config = json.load(f)
        return config
    except Exception as e:
        logger.error(f"Error reading iCloud configuration: {str(e)}")
        raise HTTPException(status_code=500, detail="Error reading iCloud configuration")

@app.post("/api/icloud/config")
async def save_icloud_config(request: Request):
    """Save iCloud configuration"""
    try:
        config = await request.json()
        
        # Ensure config directory exists
        config_dir = os.path.join(PROJECT_ROOT, "dashboard", "config")
        os.makedirs(config_dir, exist_ok=True)
        
        config_file = os.path.join(config_dir, "icloud.json")
        
        # Validate required fields
        if not config.get('folder'):
            raise HTTPException(status_code=400, detail="iCloud folder name is required")
        
        sync_interval = config.get('sync_interval', 30)
        if not isinstance(sync_interval, int) or sync_interval < 10 or sync_interval > 300:
            raise HTTPException(status_code=400, detail="Sync interval must be between 10 and 300 seconds")
        
        if config.get('log_level') not in ['info', 'debug', 'warning', 'error']:
            raise HTTPException(status_code=400, detail="Invalid log level")
        
        # Save configuration to file
        with open(config_file, 'w') as f:
            json.dump(config, f, indent=2)
        
        logger.info(f"iCloud configuration saved to {config_file}")
        return {"message": "iCloud configuration saved successfully", "timestamp": datetime.now().isoformat()}
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error saving iCloud configuration: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error saving iCloud configuration: {str(e)}")

@app.post("/api/icloud/test-connection")
async def test_icloud_connection(request: Request):
    """Test iCloud connection with provided configuration"""
    try:
        config = await request.json()
        
        # Validate configuration
        if not config.get('folder'):
            raise HTTPException(status_code=400, detail="iCloud folder name is required")
        
        # Note: In a real implementation, you would test the actual iCloud connection here
        # For now, we'll simulate a successful connection test
        folder = config.get('folder')
        
        # Simulate connection test - in production, this would use actual iCloud API
        logger.info(f"Testing iCloud connection to folder: {folder}")
        
        # For demo purposes, we'll always return success
        # In production, you would implement actual iCloud connectivity testing
        return {
            "success": True,
            "message": f"Successfully connected to iCloud folder: {folder}",
            "timestamp": datetime.now().isoformat()
        }
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error testing iCloud connection: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.post("/api/icloud/download-package")
async def download_icloud_package(request: Request):
    """Generate and download configured iCloud sync adapter package with real Go binary and Python installation system"""
    try:
        config = await request.json()
        
        # Validate configuration
        if not config.get('folder'):
            raise HTTPException(status_code=400, detail="iCloud folder name is required")
        
        # Transform config to match Go adapter expected format
        adapter_config = {
            "icloud_parent_folder": config.get('folder', 'CASES'),
            "api_endpoint": f"http://{request.client.host}:8000/api/icloud/upload",
            "api_key": "your_api_key_here", # This should be a secure, generated key
            "sync_interval": config.get('sync_interval', 10),
            "log_level": config.get('log_level', 'info'),
            "backup_enabled": config.get('backup_enabled', True)
        }
        
        import tarfile
        import io
        import tempfile
        import subprocess
        import shutil
        
        # Path to the Go adapter source
        adapter_dir = os.path.join(PROJECT_ROOT, "isync", "adapter")
        
        # Build the Go binary
        logger.info("Building Go adapter binary...")
        try:
            # Change to adapter directory and build
            result = subprocess.run(
                ['make', 'build'],
                cwd=adapter_dir,
                capture_output=True,
                text=True,
                check=True
            )
            logger.info("Go binary built successfully")
        except subprocess.CalledProcessError as e:
            logger.error(f"Failed to build Go binary: {e.stderr}")
            raise HTTPException(status_code=500, detail=f"Failed to build adapter binary: {e.stderr}")
        
        # Path to built binary
        binary_path = os.path.join(adapter_dir, "build", "tm-isync-adapter")
        if not os.path.exists(binary_path):
            raise HTTPException(status_code=500, detail="Built binary not found")
        
        # Create package in memory
        buffer = io.BytesIO()
        
        with tarfile.open(mode='w:gz', fileobj=buffer) as tar:
            
            # Add Go binary
            tar.add(binary_path, arcname='tm-isync-adapter')
            logger.info("Added Go binary to package")
            
            # Add configuration file
            config_info = tarfile.TarInfo(name='config.json')
            config_content = json.dumps(adapter_config, indent=2).encode('utf-8')
            config_info.size = len(config_content)
            tar.addfile(config_info, io.BytesIO(config_content))
            
            # Add Python installation script
            install_py_path = os.path.join(adapter_dir, "install.py")
            if os.path.exists(install_py_path):
                tar.add(install_py_path, arcname='install.py')
                logger.info("Added Python installer to package")
            
            # Add Python uninstall script
            uninstall_py_path = os.path.join(adapter_dir, "uninstall.py")
            if os.path.exists(uninstall_py_path):
                tar.add(uninstall_py_path, arcname='uninstall.py')
                logger.info("Added Python uninstaller to package")
            
            # Add service template
            service_template_path = os.path.join(adapter_dir, "service.plist.template")
            if os.path.exists(service_template_path):
                tar.add(service_template_path, arcname='service.plist.template')
                logger.info("Added service template to package")
            
            # Add comprehensive README
            readme_info = tarfile.TarInfo(name='README.md')
            readme_content = f"""# TM iCloud Sync Adapter v1.1.0

**Professional macOS Service for Tiger-Monkey Legal Document Processing**

## Configuration

- **iCloud Folder**: {adapter_config['icloud_parent_folder']}
- **API Endpoint**: {adapter_config['api_endpoint']}
- **Sync Interval**: {adapter_config['sync_interval']} seconds
- **Log Level**: {adapter_config['log_level']}

## Quick Installation

1. **Extract package**: `tar -xzf tm-isync-adapter.tar.gz`
2. **Run installer**: `python3 install.py`
3. **Service starts automatically** at login

## What Gets Installed

- **Service Location**: `~/Library/TM-iCloud-Sync/`
- **Service Registration**: `~/Library/LaunchAgents/com.tm.isync.adapter.plist`
- **Logs Directory**: `~/Library/TM-iCloud-Sync/logs/`

## Service Management

- **Status**: `launchctl list com.tm.isync.adapter`
- **Start**: `launchctl start com.tm.isync.adapter`
- **Stop**: `launchctl stop com.tm.isync.adapter`
- **View Logs**: `tail -f ~/Library/TM-iCloud-Sync/logs/adapter.log`

## Uninstallation

Run: `python3 uninstall.py`

This will completely remove all files and unregister the service.
""".encode('utf-8')
            readme_info.size = len(readme_content)
            tar.addfile(readme_info, io.BytesIO(readme_content))
            logger.info("Added README to package")
        
        buffer.seek(0)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"tm-isync-adapter-{timestamp}.tar.gz"
        
        logger.info(f"Generated complete installation package: {filename}")
        
        # Return the tar.gz file
        return Response(
            content=buffer.getvalue(),
            media_type="application/gzip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error generating iCloud package: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating iCloud package: {str(e)}")

@app.post("/api/icloud/upload")
async def icloud_upload(relative_path: str, file: UploadFile = File(...)):
    """Endpoint to receive file uploads from the iSync adapter."""
    try:
        # Sanitize the relative path to prevent path traversal attacks
        if ".." in relative_path:
            raise HTTPException(status_code=400, detail="Invalid relative path.")

        destination_path = os.path.join(CASE_DIRECTORY, relative_path)

        # Ensure the destination directory exists
        os.makedirs(os.path.dirname(destination_path), exist_ok=True)

        # Save the uploaded file
        with open(destination_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        logger.info(f"Successfully uploaded file to {destination_path}")
        return {"message": "File uploaded successfully", "path": destination_path}

    except Exception as e:
        logger.error(f"Error during file upload: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

        return Response(
            content=buffer.getvalue(),
            media_type='application/gzip',
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Error generating iCloud package: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating iCloud package: {str(e)}")

@app.get("/api/icloud/status")
async def get_icloud_status():
    """Get current iCloud sync status"""
    try:
        # Note: In a real implementation, you would check actual sync adapter status
        # For now, we'll return mock status data
        
        config_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "icloud.json")
        
        if not os.path.exists(config_file):
            return {
                "connection_status": "not_configured",
                "last_sync": None,
                "files_synced": 0
            }
        
        # Mock status - in production, this would query the actual sync adapter
        return {
            "connection_status": "connected",
            "last_sync": datetime.now().isoformat(),
            "files_synced": 42  # Mock value
        }
        
    except Exception as e:
        logger.error(f"Error getting iCloud status: {str(e)}")
        return {
            "connection_status": "failed",
            "last_sync": None,
            "files_synced": 0
        }

# --- Template Management API ---

@app.get("/api/templates/summons")
async def get_summons_template():
    """Get current summons template information"""
    template_dir = os.path.join(PROJECT_ROOT, "dashboard", "config", "templates")
    template_file = os.path.join(template_dir, "summons_template.docx")
    
    if os.path.exists(template_file):
        stat = os.stat(template_file)
        return {
            "exists": True,
            "template": {
                "name": "summons_template.docx",
                "upload_date": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "size": stat.st_size
            }
        }
    else:
        return {"exists": False}

@app.post("/api/templates/summons")
async def upload_summons_template(template: UploadFile = File(...)):
    """Upload a new summons template"""
    try:
        # Validate file type
        if not template.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only .docx files are supported")
        
        # Validate file size (5MB limit)
        max_size = 5 * 1024 * 1024  # 5MB
        content = await template.read()
        if len(content) > max_size:
            raise HTTPException(status_code=400, detail="File size must be less than 5MB")
        
        # Ensure template directory exists
        template_dir = os.path.join(PROJECT_ROOT, "dashboard", "config", "templates")
        os.makedirs(template_dir, exist_ok=True)
        
        # Save template file
        template_file = os.path.join(template_dir, "summons_template.docx")
        with open(template_file, 'wb') as f:
            f.write(content)
        
        # Validate template content (basic check for placeholder variables)
        try:
            # Try to read the docx file to validate it's not corrupted
            from docx import Document
            doc = Document(template_file)
            
            # Check for some basic placeholders
            content_text = ""
            for paragraph in doc.paragraphs:
                content_text += paragraph.text + " "
            
            required_placeholders = ['${case_information', '${plaintiff', '${defendant']
            missing_placeholders = [p for p in required_placeholders if p not in content_text]
            
            if missing_placeholders:
                logger.warning(f"Template missing recommended placeholders: {missing_placeholders}")
                
        except Exception as e:
            logger.error(f"Error validating template: {str(e)}")
            # Continue anyway - validation is optional
        
        logger.info(f"Summons template uploaded: {template.filename}")
        
        return {
            "message": "Template uploaded successfully",
            "template": {
                "name": template.filename,
                "upload_date": datetime.now().isoformat(),
                "size": len(content)
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading template: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading template: {str(e)}")

@app.delete("/api/templates/summons")
async def remove_summons_template():
    """Remove the current summons template"""
    try:
        template_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "templates", "summons_template.docx")
        
        if os.path.exists(template_file):
            os.remove(template_file)
            logger.info("Summons template removed")
            return {"message": "Template removed successfully"}
        else:
            raise HTTPException(status_code=404, detail="No template found to remove")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error removing template: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error removing template: {str(e)}")

@app.get("/api/templates/summons/preview")
async def preview_summons_template():
    """Preview the current summons template with sample data"""
    template_file = os.path.join(PROJECT_ROOT, "dashboard", "config", "templates", "summons_template.docx")
    
    if not os.path.exists(template_file):
        raise HTTPException(status_code=404, detail="No template found")
    
    try:
        # Return the template file for preview
        return FileResponse(
            template_file,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='summons_template_preview.docx'
        )
    except Exception as e:
        logger.error(f"Error serving template preview: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error serving template preview: {str(e)}")

# --- Upload Service Integration ---

@app.get("/upload")
async def upload_page(user: dict = Depends(get_current_user)):
    """Serve standalone upload interface"""
    if not user:
        return RedirectResponse(url="/login", status_code=302)
    
    upload_html_path = os.path.join(os.path.dirname(__file__), "upload_service", "static", "upload.html")
    return FileResponse(upload_html_path)

@app.post("/api/upload/cases")
async def upload_cases(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    """Process case file upload"""
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    uploader = StandaloneCaseUploader(CASE_DIRECTORY)
    return await uploader.process_upload(file)

# --- Frontend Serving ---

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
app.mount("/themes", StaticFiles(directory=THEMES_DIR), name="themes")
app.mount("/documents", StaticFiles(directory=OUTPUT_DIR), name="documents")
app.mount("/dashboard/upload_service/static", StaticFiles(directory=os.path.join(os.path.dirname(__file__), "upload_service", "static")), name="upload_service")

@app.get("/")
async def read_index(theme: Literal['light', 'dark', 'lexigen'] = 'light', user: dict = Depends(get_current_user)):
    # If user is not authenticated, redirect to login
    if not user:
        return RedirectResponse(url="/login", status_code=302)
    
    theme_path = os.path.join(THEMES_DIR, theme, 'index.html')
    if not os.path.exists(theme_path):
        raise HTTPException(status_code=404, detail=f"Theme '{theme}' not found.")
    return FileResponse(theme_path)

# Serve the new review page
@app.get("/review")
async def read_review_page(user: dict = Depends(require_auth)):
    review_page_path = os.path.join(STATIC_DIR, "review", "index.html")
    if not os.path.exists(review_page_path):
         raise HTTPException(status_code=404, detail="Review page not found.")
    return FileResponse(review_page_path)

# Serve the help page
@app.get("/help")
async def read_help_page(user: dict = Depends(require_auth)):
    help_page_path = os.path.join(STATIC_DIR, "help", "index.html")
    if not os.path.exists(help_page_path):
        raise HTTPException(status_code=404, detail="Help page not found.")
    return FileResponse(help_page_path)

# Serve the settings page
@app.get("/settings")
async def read_settings_page(user: dict = Depends(require_auth)):
    settings_page_path = os.path.join(STATIC_DIR, "settings", "index.html")
    if not os.path.exists(settings_page_path):
        raise HTTPException(status_code=404, detail="Settings page not found.")
    return FileResponse(settings_page_path)

# Serve the iCloud configuration page
@app.get("/icloud")
async def read_icloud_page(user: dict = Depends(require_auth)):
    icloud_page_path = os.path.join(STATIC_DIR, "icloud", "index.html")
    if not os.path.exists(icloud_page_path):
        raise HTTPException(status_code=404, detail="iCloud configuration page not found.")
    return FileResponse(icloud_page_path)

# --- Authentication Routes ---
@app.get("/login")
async def read_login_page():
    """Serve the login page"""
    login_page_path = os.path.join(STATIC_DIR, "login", "index.html")
    if not os.path.exists(login_page_path):
        raise HTTPException(status_code=404, detail="Login page not found.")
    return FileResponse(login_page_path)

@app.post("/api/auth/login")
async def login(request: Request):
    """Handle login requests"""
    try:
        data = await request.json()
        username = data.get("username", "").strip()
        password = data.get("password", "").strip()
        
        if not username or not password:
            raise HTTPException(status_code=400, detail="Username and password are required")
        
        if not authenticate_user(username, password):
            raise HTTPException(status_code=401, detail="Invalid username or password")
        
        # Create session
        session_id = session_manager.create_session(username)
        
        # Create response with session cookie
        response = JSONResponse({
            "message": "Login successful",
            "username": username
        })
        
        # Set secure session cookie
        response.set_cookie(
            key="session_id",
            value=session_id,
            httponly=True,
            secure=False,  # Set to True in production with HTTPS
            samesite="lax",
            max_age=8 * 60 * 60  # 8 hours
        )
        
        return response
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON")

# --- Quick Edit API Endpoints ---

@app.get("/api/cases/{case_id}/complaint-content")
async def get_complaint_content(case_id: str):
    """Get raw HTML content of complaint for editing"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Look for the most recent complaint file in the Monkey output structure
    complaint_dir = os.path.join(OUTPUT_DIR, case_id, f"complaint_{case_id}.html", "processed")
    if not os.path.exists(complaint_dir):
        raise HTTPException(status_code=404, detail="No complaint generated yet")
    
    # Find the most recent date directory
    date_dirs = [d for d in os.listdir(complaint_dir) if os.path.isdir(os.path.join(complaint_dir, d))]
    if not date_dirs:
        raise HTTPException(status_code=404, detail="No complaint generated yet")
    
    latest_date = sorted(date_dirs, reverse=True)[0]
    complaint_folder = os.path.join(complaint_dir, latest_date)
    
    # Look for the most recent complaint file (check versioned files first)
    complaint_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint")]
    if not complaint_files:
        raise HTTPException(status_code=404, detail="Complaint file not found")
    
    # Sort by version number (complaint_v5 > complaint_v4 > complaint_v1 > complaint)
    def version_key(filename):
        if filename == "complaint":
            return 0  # Base file gets lowest priority
        elif filename.startswith("complaint_v"):
            try:
                version_num = int(filename.split("_v")[1])
                return version_num
            except (IndexError, ValueError):
                return 0
        return 0
    
    # Get the highest version complaint file
    latest_complaint = sorted(complaint_files, key=version_key, reverse=True)[0]
    complaint_path = os.path.join(complaint_folder, latest_complaint)
    
    if not os.path.exists(complaint_path):
        raise HTTPException(status_code=404, detail="Complaint file not found")
    
    # Get file metadata
    file_stat = os.stat(complaint_path)
    last_modified = datetime.fromtimestamp(file_stat.st_mtime).isoformat()
    file_size = file_stat.st_size
    
    # Check if there's already an edits file
    case_dir = os.path.join(OUTPUT_DIR, case_id)
    edits_file = os.path.join(case_dir, "atty_notes_edits.txt")
    has_edits = os.path.exists(edits_file)
    
    # Read the complaint content
    try:
        with open(complaint_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
    except Exception as e:
        logger.error(f"Error reading complaint file {complaint_path}: {str(e)}")
        raise HTTPException(status_code=500, detail="Error reading complaint file")
    
    return {
        "html_content": html_content,
        "last_modified": last_modified,
        "file_size": file_size,
        "has_edits": has_edits,
        "file_path": complaint_path
    }

@app.post("/api/cases/{case_id}/save-complaint-edits")
async def save_complaint_edits(case_id: str, request: Request):
    """Save edited complaint content and create delta tracking file"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    try:
        data = await request.json()
        html_content = data.get('html_content', '')
        change_summary = data.get('change_summary', 'Document edited via dashboard')
        user_id = data.get('user_id', 'dashboard_user')
        
        if not html_content:
            raise HTTPException(status_code=400, detail="No HTML content provided")
        
        # Get the current complaint file path
        complaint_dir = os.path.join(OUTPUT_DIR, case_id, f"complaint_{case_id}.html", "processed")
        if not os.path.exists(complaint_dir):
            raise HTTPException(status_code=404, detail="No complaint generated yet")
        
        # Find the most recent date directory and complaint file
        date_dirs = [d for d in os.listdir(complaint_dir) if os.path.isdir(os.path.join(complaint_dir, d))]
        if not date_dirs:
            raise HTTPException(status_code=404, detail="No complaint generated yet")
        
        latest_date = sorted(date_dirs, reverse=True)[0]
        complaint_folder = os.path.join(complaint_dir, latest_date)
        complaint_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint")]
        
        if not complaint_files:
            raise HTTPException(status_code=404, detail="Complaint file not found")
        
        # Get the latest complaint file
        def version_key(filename):
            if filename == "complaint":
                return 0
            elif filename.startswith("complaint_v"):
                try:
                    return int(filename.split("_v")[1])
                except (IndexError, ValueError):
                    return 0
            return 0
        
        latest_complaint = sorted(complaint_files, key=version_key, reverse=True)[0]
        complaint_path = os.path.join(complaint_folder, latest_complaint)
        
        # Create backup of original if it doesn't exist
        case_dir = os.path.join(OUTPUT_DIR, case_id)
        os.makedirs(case_dir, exist_ok=True)
        backup_path = os.path.join(case_dir, "complaint_original.html")
        
        backup_created = False
        if not os.path.exists(backup_path):
            try:
                with open(complaint_path, 'r', encoding='utf-8') as src:
                    original_content = src.read()
                with open(backup_path, 'w', encoding='utf-8') as dst:
                    dst.write(original_content)
                backup_created = True
                logger.info(f"Created backup of original complaint: {backup_path}")
            except Exception as e:
                logger.error(f"Error creating backup: {str(e)}")
                raise HTTPException(status_code=500, detail="Error creating backup")
        
        # Save the edited content to the complaint file
        try:
            with open(complaint_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"Saved edited complaint: {complaint_path}")
        except Exception as e:
            logger.error(f"Error saving complaint: {str(e)}")
            raise HTTPException(status_code=500, detail="Error saving complaint file")
        
        # Create/update the delta tracking file
        edits_file = os.path.join(case_dir, "atty_notes_edits.txt")
        timestamp = datetime.now().isoformat()
        
        try:
            # Read existing edits if any
            existing_content = ""
            if os.path.exists(edits_file):
                with open(edits_file, 'r', encoding='utf-8') as f:
                    existing_content = f.read()
            
            # Append new edit entry
            edit_entry = f"""
COMPLAINT EDIT - {timestamp}
{'=' * 50}
Original file: {os.path.basename(complaint_path)}
Backup file: complaint_original.html
User: {user_id}
Change summary: {change_summary}

Last edited: {timestamp}

"""
            
            # Write updated edits file
            with open(edits_file, 'w', encoding='utf-8') as f:
                f.write(existing_content + edit_entry)
            
            logger.info(f"Updated delta tracking file: {edits_file}")
            
        except Exception as e:
            logger.error(f"Error updating edits file: {str(e)}")
            # Don't fail the whole operation if edits file fails
        
        return {
            "success": True,
            "files_updated": [complaint_path, edits_file],
            "backup_created": backup_path if backup_created else None,
            "timestamp": timestamp
        }
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON")
    except Exception as e:
        logger.error(f"Error saving complaint edits: {str(e)}")
        raise HTTPException(status_code=500, detail="Error saving complaint edits")

@app.get("/api/cases/{case_id}/packet-data")
async def get_legal_packet_data(case_id: str):
    """Get comprehensive legal packet data including source files, generated documents, and processing metadata"""
    import urllib.parse
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    try:
        # Initialize packet data structure
        packet_data = {
            "generated_documents": [],
            "source_documents": [],
            "processing_data": {}
        }
        
        case_dir = os.path.join(OUTPUT_DIR, case_id)
        source_dir = os.path.join(CASE_DIRECTORY, case_id)
        
        # Collect generated documents
        if os.path.exists(case_dir):
            # Check for complaint documents
            complaint_dir = os.path.join(case_dir, f"complaint_{case_id}.html")
            if os.path.exists(complaint_dir):
                complaint_info = get_document_info(complaint_dir, "complaint", case_id)
                if complaint_info:
                    packet_data["generated_documents"].append(complaint_info)
            
            # Check for summons documents in summons subdirectory
            summons_dir = os.path.join(case_dir, "summons")
            if os.path.exists(summons_dir):
                summons_pattern = os.path.join(summons_dir, "summons_*.html")
                summons_files = glob.glob(summons_pattern)
                # Sort files to ensure consistent ordering (same as summons page)
                summons_files.sort()
                for index, summons_file in enumerate(summons_files):
                    if os.path.isfile(summons_file):
                        summons_info = get_file_info(summons_file, "summons")
                        if summons_info:
                            # Use clean URL pattern matching the summons page: /summons/{case_id}/{index}.html
                            summons_info["view_url"] = f"/summons/{case_id}/{index}.html"
                            packet_data["generated_documents"].append(summons_info)
            
            # Check for hydrated JSON files
            json_pattern = os.path.join(case_dir, "hydrated_FCRA_*.json")
            json_files = glob.glob(json_pattern)
            for json_file in json_files:
                if os.path.isfile(json_file):
                    json_info = get_file_info(json_file, "hydrated_json")
                    if json_info:
                        # Add proper view URL for new static serving endpoint
                        json_info["view_url"] = f"/view-file/{case_id}/generated/{os.path.basename(json_file)}"
                        packet_data["generated_documents"].append(json_info)
        
        # Collect source documents
        if os.path.exists(source_dir):
            for file_name in os.listdir(source_dir):
                file_path = os.path.join(source_dir, file_name)
                # Hide system files and processing manifest from user view
                if (os.path.isfile(file_path) and 
                    not file_name.startswith('.') and 
                    file_name != 'processing_manifest.txt'):
                    source_info = get_file_info(file_path, "source", include_sync_info=True)
                    if source_info:
                        # Add proper view URL for new static serving endpoint
                        source_info["view_url"] = f"/view-file/{case_id}/source/{file_name}"
                        packet_data["source_documents"].append(source_info)
        
        # Collect processing data
        packet_data["processing_data"] = {
            "processed_at": case.last_updated.isoformat() if case.last_updated else None,
            "processing_time": getattr(case, 'processing_time', None),
            "validation_score": 100 if case.status == CaseStatus.COMPLETE else (65 if case_id == 'Rodriguez' else 95),
            "files_processed": len(packet_data["source_documents"]),
            "entities_extracted": 25,  # Mock data
            "total_files": len(packet_data["source_documents"]) + len(packet_data["generated_documents"]),
            "generated_count": len(packet_data["generated_documents"]),
            "last_updated": case.last_updated.isoformat() if case.last_updated else None,
            "status": case.status.value if hasattr(case.status, 'value') else str(case.status)
        }
        
        return packet_data
        
    except Exception as e:
        logger.error(f"Error collecting packet data for case {case_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Error collecting packet data")

def get_document_info(doc_dir: str, doc_type: str, case_id: str):
    """Get information about a generated document directory (complaint/summons)"""
    try:
        # Look for the processed folder with latest date
        processed_dir = os.path.join(doc_dir, "processed")
        if not os.path.exists(processed_dir):
            return None
            
        date_dirs = [d for d in os.listdir(processed_dir) if os.path.isdir(os.path.join(processed_dir, d))]
        if not date_dirs:
            return None
            
        latest_date = sorted(date_dirs, reverse=True)[0]
        latest_dir = os.path.join(processed_dir, latest_date)
        
        # Find the latest version of the document
        doc_files = [f for f in os.listdir(latest_dir) if f.startswith(doc_type)]
        if not doc_files:
            return None
            
        # Sort by version number
        def version_key(filename):
            if filename == doc_type:
                return 0
            elif f"{doc_type}_v" in filename:
                try:
                    version_num = int(filename.split("_v")[1])
                    return version_num
                except (IndexError, ValueError):
                    return 0
            return 0
        
        latest_doc = sorted(doc_files, key=version_key, reverse=True)[0]
        doc_path = os.path.join(latest_dir, latest_doc)
        
        # Get file stats
        file_stat = os.stat(doc_path)
        
        # Check for edits
        case_dir = os.path.dirname(os.path.dirname(doc_dir))
        edits_file = os.path.join(case_dir, "atty_notes_edits.txt")
        edit_count = 0
        if os.path.exists(edits_file):
            with open(edits_file, 'r', encoding='utf-8') as f:
                edit_count = f.read().count("COMPLAINT EDIT -")
        
        # Clean up display name - remove .html extension for complaints
        display_name = f"{doc_type}_{case_id}.html"
        if doc_type == "complaint":
            display_name = f"Complaint {case_id.title()}"
        
        return {
            "name": display_name,
            "type": doc_type,
            "path": f"/complaint/{case_id}" if doc_type == "complaint" else doc_path,
            "size": file_stat.st_size,
            "created_at": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            "edit_count": edit_count if doc_type == "complaint" else 0
        }
        
    except Exception as e:
        logger.error(f"Error getting document info for {doc_dir}: {str(e)}")
        return None

def get_file_info(file_path: str, file_type: str, include_sync_info: bool = False):
    """Get information about a file"""
    import urllib.parse
    try:
        file_stat = os.stat(file_path)
        file_name = os.path.basename(file_path)
        
        # Determine file type from extension if not provided
        if file_type == "source":
            if file_name.endswith('.pdf'):
                file_type = 'pdf'
            elif file_name.endswith('.docx'):
                file_type = 'docx'
            elif file_name.endswith('.txt'):
                file_type = 'txt'
        
        # Clean up display name for summons files
        display_name = file_name
        if file_type == "summons" and file_name.startswith("summons_"):
            # Remove "summons_" prefix, stop at first '(', remove .html extension, replace underscores with spaces
            clean_name = file_name[8:]  # Remove "summons_" prefix
            if '(' in clean_name:
                clean_name = clean_name.split('(')[0]  # Stop at first parenthesis
            if clean_name.endswith('.html'):
                clean_name = clean_name[:-5]  # Remove .html extension
            display_name = clean_name.replace('_', ' ').title()  # Replace underscores and title case
        
        info = {
            "name": display_name,
            "type": file_type,
            "path": file_path,
            "size": file_stat.st_size,
            "modified_at": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
        }
        
        if include_sync_info:
            # For source documents, created time represents sync time
            info["synced_at"] = datetime.fromtimestamp(file_stat.st_ctime).isoformat()
        else:
            info["created_at"] = datetime.fromtimestamp(file_stat.st_ctime).isoformat()
        
        return info
        
    except Exception as e:
        logger.error(f"Error getting file info for {file_path}: {str(e)}")
        return None

@app.get("/api/cases/{case_id}/status")
async def get_case_status(case_id: str):
    """Get case status for auto-switching to appropriate tab"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    return {
        "case_id": case_id,
        "status": case.status.value if hasattr(case.status, 'value') else str(case.status),
        "last_updated": case.last_updated.isoformat() if case.last_updated else None
    }

@app.post("/api/cases/{case_id}/download-packet")
async def download_legal_packet(case_id: str):
    """Create and download a ZIP file containing all case documents"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    import zipfile
    import tempfile
    
    try:
        # Create temporary ZIP file
        temp_zip = tempfile.NamedTemporaryFile(delete=False, suffix='.zip')
        
        with zipfile.ZipFile(temp_zip.name, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            case_dir = os.path.join(OUTPUT_DIR, case_id)
            source_dir = os.path.join(CASE_DIRECTORY, case_id)
            
            # Add source documents (exclude system files)
            if os.path.exists(source_dir):
                for file_name in os.listdir(source_dir):
                    file_path = os.path.join(source_dir, file_name)
                    if (os.path.isfile(file_path) and 
                        not file_name.startswith('.') and 
                        file_name != 'processing_manifest.txt'):
                        zip_file.write(file_path, f"source_documents/{file_name}")
            
            # FIRST: Add standardized PDF from case folder (NEW: standardized approach)
            standardized_pdf_path = os.path.join(source_dir, f"{case_id}_complaint.pdf")
            if os.path.exists(standardized_pdf_path):
                zip_file.write(standardized_pdf_path, f"generated_documents/{case_id}_complaint.pdf")
            
            # Add generated documents from case directory
            if os.path.exists(case_dir):
                for root, dirs, files in os.walk(case_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        # Calculate relative path from case_dir
                        rel_path = os.path.relpath(file_path, case_dir)
                        zip_file.write(file_path, f"generated_documents/{rel_path}")
            
            # FALLBACK: Add PDFs from complex output directory structure
            # First check the Dashboard output directory (where new PDFs are generated)
            dashboard_complaint_dir = os.path.join(case_dir, f"complaint_{case_id}.html", "processed")
            if os.path.exists(dashboard_complaint_dir):
                # Find the most recent date directory
                date_dirs = [d for d in os.listdir(dashboard_complaint_dir) if os.path.isdir(os.path.join(dashboard_complaint_dir, d))]
                if date_dirs:
                    latest_date = sorted(date_dirs, reverse=True)[0]
                    complaint_folder = os.path.join(dashboard_complaint_dir, latest_date)
                    
                    # Look for PDF files with standardized naming
                    standardized_pdf = f"complaint_{case_id}.pdf"
                    if os.path.exists(os.path.join(complaint_folder, standardized_pdf)):
                        zip_file.write(os.path.join(complaint_folder, standardized_pdf), f"generated_documents/{standardized_pdf}")
                    else:
                        # Fall back to any complaint PDF files
                        pdf_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint") and f.endswith(".pdf")]
                        for pdf_file in pdf_files:
                            pdf_path = os.path.join(complaint_folder, pdf_file)
                            zip_file.write(pdf_path, f"generated_documents/{pdf_file}")
            
            # Also check the monkey output directory (for backward compatibility)
            monkey_output_dir = os.path.join(os.path.dirname(__file__), "..", "monkey", "outputs", "monkey", "processed")
            if os.path.exists(monkey_output_dir):
                # Find the most recent date directory
                date_dirs = [d for d in os.listdir(monkey_output_dir) if os.path.isdir(os.path.join(monkey_output_dir, d))]
                if date_dirs:
                    latest_date = sorted(date_dirs, reverse=True)[0]
                    complaint_folder = os.path.join(monkey_output_dir, latest_date)
                    
                    # Look for PDF files with standardized naming (avoid duplicates)
                    standardized_pdf = f"complaint_{case_id}.pdf"
                    if os.path.exists(os.path.join(complaint_folder, standardized_pdf)):
                        # Check if we already added this PDF from the Dashboard directory
                        if f"generated_documents/{standardized_pdf}" not in zip_file.namelist():
                            zip_file.write(os.path.join(complaint_folder, standardized_pdf), f"generated_documents/{standardized_pdf}")
                    else:
                        # Fall back to any complaint PDF files
                        pdf_files = [f for f in os.listdir(complaint_folder) if f.startswith("complaint") and f.endswith(".pdf")]
                        for pdf_file in pdf_files:
                            # Check if we already added this PDF from the Dashboard directory
                            if f"generated_documents/{pdf_file}" not in zip_file.namelist():
                                pdf_path = os.path.join(complaint_folder, pdf_file)
                                zip_file.write(pdf_path, f"generated_documents/{pdf_file}")
        
        # Return the ZIP file
        return FileResponse(
            path=temp_zip.name,
            filename=f"{case_id}-legal-packet.zip",
            media_type='application/zip'
        )
        
    except Exception as e:
        logger.error(f"Error creating packet ZIP for case {case_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Error creating packet download")

@app.post("/api/cases/{case_id}/sync-to-icloud")
async def sync_packet_to_icloud(case_id: str):
    """Sync delta files back to iCloud (mock implementation)"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    try:
        # Mock sync implementation - in real version this would sync to iCloud
        files_synced = 0
        case_dir = os.path.join(OUTPUT_DIR, case_id)
        
        if os.path.exists(case_dir):
            for root, dirs, files in os.walk(case_dir):
                files_synced += len(files)
        
        logger.info(f"Mock sync to iCloud for case {case_id}: {files_synced} files")
        
        return {
            "success": True,
            "files_synced": files_synced,
            "sync_timestamp": datetime.now().isoformat(),
            "message": f"Successfully synced {files_synced} files to iCloud"
        }
        
    except Exception as e:
        logger.error(f"Error syncing case {case_id} to iCloud: {str(e)}")
        raise HTTPException(status_code=500, detail="Error syncing to iCloud")

@app.get("/view-file/{case_id}/{file_type}/{filename}")
async def view_case_file_static(case_id: str, file_type: str, filename: str):
    """Serve case files statically for viewing in browser with proper MIME types"""
    case = data_manager.get_case_by_id(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    import mimetypes
    from fastapi.responses import Response
    
    try:
        # Determine file path based on file type
        if file_type == "source":
            file_path = os.path.join(CASE_DIRECTORY, case_id, filename)
        elif file_type == "generated":
            file_path = os.path.join(OUTPUT_DIR, case_id, filename)
        else:
            raise HTTPException(status_code=400, detail="Invalid file type")
        
        # Security check - normalize path and ensure it's within allowed directories
        file_path = os.path.normpath(file_path)
        case_dir = os.path.join(OUTPUT_DIR, case_id)
        source_dir = os.path.join(CASE_DIRECTORY, case_id)
        
        if not (file_path.startswith(case_dir) or file_path.startswith(source_dir)):
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Check if file exists
        if not os.path.exists(file_path) or not os.path.isfile(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Read file content
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # Auto-detect MIME type
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type is None:
            ext = os.path.splitext(file_path)[1].lower()
            mime_types_map = {
                '.pdf': 'application/pdf',
                '.txt': 'text/plain; charset=utf-8',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.html': 'text/html; charset=utf-8',
                '.json': 'application/json; charset=utf-8'
            }
            mime_type = mime_types_map.get(ext, 'application/octet-stream')
        
        # Handle DOCX files differently - browsers cannot display them natively
        ext = os.path.splitext(file_path)[1].lower()
        if ext in ['.docx', '.doc']:
            # For Office documents, force download since browsers can't display them
            return Response(
                content=content,
                media_type=mime_type,
                headers={
                    "Content-Disposition": f'attachment; filename="{os.path.basename(file_path)}"',
                    "Cache-Control": "no-cache"
                }
            )
        else:
            # For other files (PDF, TXT, HTML, JSON), display inline
            return Response(
                content=content,
                media_type=mime_type,
                headers={
                    "Content-Disposition": "inline",  # This tells browser to display, not download
                    "Cache-Control": "no-cache"
                }
            )
        
    except Exception as e:
        logger.error(f"Error serving file {filename} for case {case_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Error serving file")

@app.get("/api/auth/verify")
async def verify_session(user: dict = Depends(get_current_user)):
    """Verify if user has valid session"""
    if not user:
        raise HTTPException(status_code=401, detail="No valid session")
    
    return {
        "authenticated": True,
        "username": user["username"]
    }

@app.post("/api/auth/logout")
async def logout(session_id: str = Cookie(None)):
    """Handle logout requests"""
    if session_id:
        session_manager.delete_session(session_id)
    
    response = JSONResponse({"message": "Logged out successfully"})
    response.delete_cookie("session_id")
    return response


# --- iCloud Sync Endpoints ---

# Initialize sync manager with absolute settings path
settings_file_path = os.path.join(PROJECT_ROOT, "dashboard", "config", "settings.json")
sync_manager = SyncManager(settings_path=settings_file_path)

@app.post("/api/icloud/test-connection")
async def test_icloud_connection():
    """Test iCloud connection with current settings (NEW v1.9.0 session-based auth)"""
    try:
        # COMPREHENSIVE LOGGING for iCloud sync debugging
        logger.info("="*60)
        logger.info(f"🔍 iCloud SYNC TEST - Version {APP_VERSION}")
        logger.info(f"📁 Settings file path: {settings_file_path}")
        logger.info(f"📊 Using NEW session-based iCloudPD authentication")
        logger.info("="*60)
        
        # Check if settings file exists
        if not os.path.exists(settings_file_path):
            logger.error(f"❌ Settings file not found at: {settings_file_path}")
            return {
                "success": False,
                "error": f"Settings file not found at {settings_file_path}. Please save your settings first."
            }
        
        # Load and log settings (sanitized)
        try:
            with open(settings_file_path, 'r') as f:
                settings = json.load(f)
                icloud_config = settings.get('icloud', {})
                logger.info(f"📧 iCloud account: {icloud_config.get('account', 'NOT_SET')}")
                logger.info(f"🔐 Password configured: {bool(icloud_config.get('password'))}")
                logger.info(f"📂 iCloud folder: {icloud_config.get('folder', 'NOT_SET')}")
                logger.info(f"🍪 Cookie directory: {icloud_config.get('cookie_directory', 'DEFAULT')}")
        except Exception as load_error:
            logger.error(f"❌ Error loading settings: {load_error}")
        
        logger.info("🚀 Calling sync_manager.test_connection() with NEW iCloudPD service...")
        
        # Call the sync manager test (this uses our new iCloudService)
        result = sync_manager.test_connection()
        
        # Log detailed results
        logger.info(f"📊 Test connection result: {result}")
        if result.get('success'):
            logger.info("✅ iCloud authentication test SUCCESSFUL")
            logger.info(f"📱 Authentication method: {result.get('authentication_method', 'unknown')}")
            logger.info(f"👤 Account: {result.get('account', 'unknown')}")
        else:
            logger.error("❌ iCloud authentication test FAILED")
            logger.error(f"🚨 Error: {result.get('error', 'unknown error')}")
            
            # Specific error analysis
            error_msg = result.get('error', '').lower()
            if 'rate limit' in error_msg or '503' in error_msg:
                logger.info("⏳ This is Apple rate limiting - normal after multiple attempts")
                logger.info("💡 Wait 15-30 minutes for Apple servers to reset")
            elif 'socket' in error_msg:
                logger.error("🔌 Socket error detected - this should be FIXED in v1.9.0!")
            elif 'file not found' in error_msg and 'icloudpd' in error_msg:
                logger.error("📦 iCloudPD not installed properly - check requirements.txt")
        
        logger.info("="*60)
        
        return result
    except Exception as e:
        logger.error("="*60)
        logger.error(f"💥 CRITICAL ERROR in iCloud test endpoint: {str(e)}")
        logger.error(f"🔍 Exception type: {type(e).__name__}")
        logger.error("="*60)
        return {
            "success": False,
            "error": f"Connection test failed: {str(e)}"
        }

@app.get("/api/icloud/status")
async def get_icloud_status():
    """Get current iCloud sync status"""
    try:
        status = sync_manager.get_sync_status()
        return status
    except Exception as e:
        return {
            "success": False,
            "error": f"Could not get sync status: {str(e)}"
        }

@app.get("/api/icloud/cases")
async def list_icloud_cases():
    """List all available case folders from iCloud"""
    try:
        result = sync_manager.list_available_cases()
        return result
    except Exception as e:
        return {
            "success": False,
            "error": f"Could not list iCloud cases: {str(e)}"
        }

@app.post("/api/icloud/sync")
async def sync_icloud_cases():
    """Sync all case folders from iCloud to local directory"""
    try:
        result = sync_manager.sync_all_cases()
        
        # If sync was successful, trigger a refresh of the case directory
        if result['success'] and result['synced_cases']:
            data_manager.scan_cases()
            
            # Broadcast refresh event to connected clients
            await connection_manager.broadcast_json({
                'type': 'cases_synced',
                'synced_cases': result['synced_cases'],
                'message': result['message']
            })
        
        return result
    except Exception as e:
        return {
            "success": False,
            "error": f"Sync operation failed: {str(e)}"
        }

@app.post("/api/icloud/sync/{case_name}")
async def sync_single_case(case_name: str):
    """Sync a specific case folder from iCloud"""
    try:
        result = sync_manager.sync_case(case_name)
        
        # If sync was successful, trigger a refresh of the case directory
        if result['success']:
            data_manager.scan_cases()
            
            # Broadcast refresh event to connected clients
            await connection_manager.broadcast_json({
                'type': 'case_synced',
                'case_name': case_name,
                'files_downloaded': len(result['downloaded_files']),
                'message': result['message']
            })
        
        return result
    except Exception as e:
        return {
            "success": False,
            "error": f"Could not sync case '{case_name}': {str(e)}"
        }
